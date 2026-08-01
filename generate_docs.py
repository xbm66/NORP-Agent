# -*- coding: utf-8 -*-
"""生成 NORP Vibe Coding Agent 使用文档和插件开发说明。"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color):
    """给表格单元格设置背景色。"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def add_code_block(doc, code_text):
    """添加代码块（等宽字体、灰色背景）。"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    # 段落背景设置
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)
    return p


def add_bullet(doc, text, level=0):
    """添加要点列表。"""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.style = doc.styles['List Bullet 2']
    return p


def create_user_manual():
    """创建使用文档 .docx (7k~1万字)。"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 样式 ──
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = '微软雅黑'
        heading_style.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    # ═══════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('NORP Vibe Coding Agent')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('使用文档')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'版本 1.0 | {datetime.date.today().strftime("%Y年%m月")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 目录页
    # ═══════════════════════════════════════════
    doc.add_heading('目  录', level=1)
    toc_items = [
        ('1', '产品简介', 3),
        ('2', '快速入门', 4),
        ('3', '界面与布局', 5),
        ('4', '会话管理', 6),
        ('5', '内置工具参考', 7),
        ('6', '配置与选项', 10),
        ('7', '插件系统', 12),
        ('8', '安全机制', 13),
        ('9', '异步架构', 14),
        ('10', '最佳实践', 15),
        ('11', '常见问题 (FAQ)', 16),
        ('12', '附录', 18),
    ]
    for num, title_text, page in toc_items:
        p = doc.add_paragraph(f'{num}.  {title_text}')
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ═══════════════════════════════
    # 1. 产品简介
    # ═══════════════════════════════
    doc.add_heading('1. 产品简介', level=1)

    doc.add_heading('1.1 什么是 NORP Vibe Coding Agent？', level=2)
    doc.add_paragraph(
        'NORP Vibe Coding Agent 是一款基于 AI 大语言模型的自主编程智能体（Coding Agent），'
        '采用 ReAct（Reasoning + Acting）架构设计。它能够将用户的自然语言指令自动转化为'
        '精确的代码操作，主动分析问题、编写代码、管理文件，而非被动地一问一答。'
    )
    doc.add_paragraph(
        '与传统的 AI 聊天机器人不同，Vibe Coding Agent 拥有完整的工具调用能力——它可以'
        '读取和修改文件、执行 Shell 命令、搜索代码、安装依赖、初始化项目，甚至联网搜索'
        '最新技术信息。它就像一个不知疲倦的全栈工程师，在你的工作区中自主完成编程任务。'
    )

    doc.add_heading('1.2 核心特性', level=2)
    features = [
        ('🤖 自主编程', '基于 ReAct 架构，自动推理并执行多步骤的编程任务，包括代码编写、调试、重构等。'),
        ('🔧 丰富的工具集', '内置 14 个工具函数，覆盖文件操作、命令执行、项目脚手架、依赖管理等场景。'),
        ('🧩 可扩展插件系统', '支持第三方插件，15 个生命周期钩子，可自定义工具和行为。多层安全审计确保插件安全。'),
        ('🛡️ 安全沙箱', '支持 Docker 容器隔离和进程组隔离，文件路径边界检查，危险命令拦截，写/删操作确认。'),
        ('📑 多会话管理', '支持最多 16 个独立会话（类似浏览器标签页），每个会话有独立的工作区和对话历史。'),
        ('⚡ 异步架构', '基于 asyncio 的异步执行引擎，沙箱池管理、文件 I/O 队列、生命周期管理。'),
        ('🔒 API 密钥加密', '使用 Windows DPAPI (win32crypt) 或系统 Keyring 加密存储 API 密钥。'),
        ('🌐 多 API 兼容', '支持 OpenAI Chat Completions、DeepSeek Responses API、Anthropic Messages API。'),
        ('💭 思维链可视化', '实时流式展示 AI 的推理过程（reasoning_content），让你看清 AI "怎么想"的。'),
        ('📊 Token 用量追踪', '实时统计 input/output/tool_call token 消耗，支持余额查询。'),
    ]
    for emoji_title, desc in features:
        p = doc.add_paragraph()
        run = p.add_run(emoji_title)
        run.bold = True
        p.add_run(f'：{desc}')

    doc.add_heading('1.3 技术架构概览', level=2)
    doc.add_paragraph(
        'Vibe Coding Agent 由以下核心模块组成：'
    )
    arch_components = [
        'main.py — 程序入口，基于 pywebview 的桌面窗口',
        'api.py — pywebview JS 桥接层，暴露所有 API 给前端',
        'async_loop.py — 异步 Agent 主循环（核心推理-行动循环）',
        'async_executor.py — 异步工具执行器（集成沙箱池、文件 I/O 队列等）',
        'config.py — 配置管理，支持加密存储 API 密钥',
        'event_queue.py — 线程安全的事件队列（Agent ↔ 前端通信）',
        'sandbox_pool.py — 沙箱池（最多 8 个沙箱，异步获取/释放）',
        'file_io_queue.py — 文件并发访问冲突检测与排队',
        'lifecycle_manager.py — 任务生命周期与僵尸进程清理',
        'permission_cascade.py — 权限级联模型（层级权限继承）',
        'resource_isolator.py — 资源隔离（终端 40% + 插件池 60%）',
        'path_mapper.py — 双向路径映射（宿主 ↔ 沙箱）',
        'plugin_system/ — 插件框架（管理器、安全审计、上下文）',
        'tools.py — 14 个内置工具的 OpenAI Function Schema 定义',
    ]
    for comp in arch_components:
        add_bullet(doc, comp)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 2. 快速入门
    # ═══════════════════════════════════
    doc.add_heading('2. 快速入门', level=1)

    doc.add_heading('2.1 安装与启动', level=2)
    doc.add_paragraph('Vibe Coding Agent 为绿色免安装桌面应用。下载后解压到任意目录，双击运行即可。')
    doc.add_paragraph('首次启动时，应用会在以下位置自动创建配置目录：')
    add_code_block(doc, '%LOCALAPPDATA%\\vibe_agent\\')
    doc.add_paragraph('该目录包含 config.json（配置文件）、base.env（加密的 API 密钥）、'
                      '插件日志、工具调用记录等。')

    doc.add_heading('2.2 配置 API 密钥', level=2)
    doc.add_paragraph('在使用 Agent 之前，需要先配置 DeepSeek API 密钥：')
    steps = [
        '打开应用，点击右上角的「设置」图标',
        '在「API 密钥」输入框中粘贴你的 DeepSeek API Key',
        '点击「验证」按钮确认密钥有效',
        '配置完成后，系统会自动加密存储密钥，无需每次输入',
    ]
    for i, step in enumerate(steps, 1):
        add_bullet(doc, f'{step}')

    doc.add_paragraph('支持的 API 端点：')
    endpoints = [
        'DeepSeek 官方端点：https://api.deepseek.com（推荐）',
        '自定义 OpenAI 兼容端点：支持任意兼容 OpenAI SDK 的 API 服务',
        'Anthropic 兼容端点：https://api.deepseek.com/anthropic',
    ]
    for ep in endpoints:
        add_bullet(doc, ep)

    doc.add_heading('2.3 第一次对话', level=2)
    doc.add_paragraph('配置好 API 密钥后，在输入框中键入你的第一个任务，例如：')
    add_code_block(doc, '帮我创建一个 Python Flask 的 Hello World 项目')
    doc.add_paragraph(
        'Agent 会自主规划步骤：初始化项目 → 创建文件 → 安装依赖 → 验证代码。'
        '你可以在界面中实时看到 AI 的思考过程和工具调用详情。'
    )

    doc.add_page_break()

    # ═══════════════════════════════════
    # 3. 界面与布局
    # ═══════════════════════════════════
    doc.add_heading('3. 界面与布局', level=1)

    doc.add_heading('3.1 主窗口', level=2)
    doc.add_paragraph('应用窗口默认大小为 1200×800 像素，最小可缩至 800×500。窗口采用现代深色主题设计，'
                      '分为以下主要区域：')
    layout = [
        ('顶部工具栏', '会话标签页切换、新建会话、设置入口、Token 用量显示'),
        ('左侧对话区', '流式展示 AI 回复内容（含推理过程折叠面板）。消息逐 token 实时渲染。'),
        ('右侧信息面板', '工具调用日志（JSON 格式）、Token 用量图表、插件状态'),
        ('底部输入区', '文本输入框 + 发送/停止按钮。支持多行输入。'),
    ]
    for name, desc in layout:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('3.2 事件流协议', level=2)
    doc.add_paragraph('前端通过轮询 EventQueue 与 Agent 通信。事件前缀含义：')
    events = [
        ('T:', '推理过程 (Thinking / Reasoning) — AI 的"内心独白"'),
        ('R:', '回复内容 (Response) — AI 的最终输出'),
        ('C:', '工具调用 (Command) — AI 决定调用某个工具'),
        ('Q:', '提问 (Question) — AI 需要用户确认或输入'),
        ('WC:', '写/删确认 (Write Confirm) — 请求用户确认文件操作'),
        ('D:', '完成 (Done) — 任务执行完毕'),
        ('E:', '错误 (Error) — 异常信息'),
        ('U:', '用量更新 (Usage) — Token 消耗统计'),
        ('F:', '思考结束 (Finalize) — 推理阶段结束，进入输出阶段'),
    ]
    for prefix, desc in events:
        p = doc.add_paragraph()
        run = p.add_run(f'{prefix} ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 4. 会话管理
    # ═══════════════════════════════════
    doc.add_heading('4. 会话管理', level=1)

    doc.add_heading('4.1 多会话架构', level=2)
    doc.add_paragraph(
        'Vibe Coding Agent 支持同时运行最多 16 个会话（类似浏览器标签页）。每个会话拥有：'
    )
    items = [
        '独立的对话历史（conversation_history）',
        '独立的事件队列（EventQueue）',
        '独立的 Agent 循环线程（AsyncAgentLoop）',
        '独立的工作区路径（workspace / project_root）',
        '独立的持久化记忆（memory.json）',
    ]
    for item in items:
        add_bullet(doc, item)

    doc.add_heading('4.2 会话操作', level=2)
    ops = [
        ('新建会话', '点击标签栏右侧的 "+" 按钮。可为新会话指定独立的工作区目录。'),
        ('切换会话', '点击标签页切换。不同会话的任务互不影响，可以并行执行。'),
        ('关闭会话', '右键标签页选择"关闭"。至少保留一个会话（最后一个无法关闭）。'),
        ('重命名会话', '双击标签页标题可编辑名称。'),
        ('设置工作区', '在会话设置中修改 project_root，该会话的所有文件操作将限定在此目录内。'),
    ]
    for name, desc in ops:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('4.3 记忆系统', level=2)
    doc.add_paragraph(
        '每个会话支持持久化记忆（Memory）。开启后，系统会自动保存对话记录到：'
    )
    add_code_block(doc, '%LOCALAPPDATA%\\vibe_agent\\memory\\memory_{session_id}.json')
    doc.add_paragraph('记忆模式有两种：')
    modes = [
        ('full（完整模式）', '保留最近 N 轮对话的完整内容。超出 max_rounds 的旧对话自动移除。'),
        ('summary（摘要模式）', '仅保留最近 2 轮对话，其余压缩为文本摘要。节省上下文 token。'),
    ]
    for name, desc in modes:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        p.add_run(f'：{desc}')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 5. 内置工具参考
    # ═══════════════════════════════════
    doc.add_heading('5. 内置工具参考', level=1)

    doc.add_paragraph(
        'Agent 拥有 14 个内置工具函数，均以 OpenAI Function Calling 格式注册。'
        '以下逐一说明每个工具的用途、参数和注意事项。'
    )

    # 5.1 read_file
    doc.add_heading('5.1 read_file — 读取文件', level=2)
    doc.add_paragraph('读取工作区内任意文本文件的内容。支持指定行范围以节省 token 消耗。')
    doc.add_paragraph('参数：')
    params = [
        'path (必填): 文件路径，相对于工作区根目录',
        'start_line (可选): 起始行号（从 1 开始）',
        'end_line (可选): 结束行号（含）',
    ]
    for p_text in params:
        add_bullet(doc, p_text)
    doc.add_paragraph('使用建议：Debug 时先用 search_in_files 定位问题行，再用行范围精准读取。')

    # 5.2 write_file
    doc.add_heading('5.2 write_file — 写入文件', level=2)
    doc.add_paragraph('创建新文件或覆盖已有文件。如果父目录不存在会自动创建。')
    doc.add_paragraph('参数：')
    add_bullet(doc, 'path (必填): 目标文件路径')
    add_bullet(doc, 'content (必填): 要写入的完整内容')
    doc.add_paragraph('⚠️ 安全提示：覆盖文件前建议先 read_file 备份原内容。'
                      '若开启了 confirm_write_delete 选项，写入前会弹出确认对话框。')

    # 5.3 replace_in_file
    doc.add_heading('5.3 replace_in_file — 精准替换', level=2)
    doc.add_paragraph('在文件中查找并替换指定文本。old_str 必须精确匹配（含缩进和换行），'
                      '且只能在文件中匹配到唯一一处。')
    doc.add_paragraph('参数：')
    add_bullet(doc, 'path (必填): 文件路径')
    add_bullet(doc, 'old_str (必填): 要被替换的原始文本（必须精确匹配）')
    add_bullet(doc, 'new_str (必填): 替换后的新文本')
    doc.add_paragraph('优势：相比 write_file 重写整个文件，此工具仅修改目标片段，大幅节省 token。')

    # 5.4 list_dir
    doc.add_heading('5.4 list_dir — 列出目录', level=2)
    doc.add_paragraph('列出指定目录下的文件和子目录。目录以 "/" 结尾标识。')
    add_bullet(doc, 'path (可选): 目录路径，默认 "." 即工作区根目录')

    # 5.5 search_in_files
    doc.add_heading('5.5 search_in_files — 搜索文件', level=2)
    doc.add_paragraph('在工作区中递归搜索包含指定文本模式的文件。自动跳过 __pycache__、node_modules、'
                      '.git 等目录。')
    add_bullet(doc, 'pattern (必填): 要搜索的文本')
    add_bullet(doc, 'path (可选): 搜索范围，可以是文件路径或目录。默认搜索整个工作区')
    doc.add_paragraph('结果限制：最多返回 50 条匹配，超出部分截断。')

    # 5.6 delete_file
    doc.add_heading('5.6 delete_file — 删除文件/目录', level=2)
    doc.add_paragraph('删除文件或整个目录（含所有子内容）。⚠️ 此操作不可逆。')
    add_bullet(doc, 'path (必填): 要删除的文件或目录路径')
    doc.add_paragraph('安全约束：执行前 Agent 必须调用 ask_user 获得用户确认。')

    # 5.7 exec_cmd
    doc.add_heading('5.7 exec_cmd — 执行命令', level=2)
    doc.add_paragraph('在沙箱或本地环境中执行 Shell 命令。内置危险命令拦截。')
    add_bullet(doc, 'command (必填): 要执行的 Shell 命令')
    add_bullet(doc, 'timeout (可选): 超时秒数，默认 30 秒')
    doc.add_paragraph('被拦截的危险模式：sudo、rm -rf /、mkfs、dd if=、format c: 等。')

    # 5.8 init_project
    doc.add_heading('5.8 init_project — 初始化项目', level=2)
    doc.add_paragraph('根据项目类型自动创建脚手架目录结构。')
    add_bullet(doc, 'type (必填): 项目类型 — python / web / node')
    add_bullet(doc, 'name (必填): 项目名称')
    doc.add_paragraph('Python 类型会创建 __init__.py、main.py、requirements.txt；'
                      'Web 类型会创建 index.html + css/ + js/ 目录。')

    # 5.9 install_dependency
    doc.add_heading('5.9 install_dependency — 安装依赖', level=2)
    doc.add_paragraph('使用 pip 或 npm 安装项目依赖。')
    add_bullet(doc, 'package (必填): 包名，如 flask、requests')
    add_bullet(doc, 'manager (可选): 包管理器 — pip（默认）、npm')

    # 5.10 git_commit
    doc.add_heading('5.10 git_commit — Git 提交', level=2)
    doc.add_paragraph('执行 git add -A 和 git commit，将当前变更提交到仓库。')
    add_bullet(doc, 'message (必填): 提交信息，建议使用约定式提交格式（如 feat: add user auth）')

    # 5.11 ask_user
    doc.add_heading('5.11 ask_user — 询问用户', level=2)
    doc.add_paragraph('当 Agent 需要用户确认、选择或补充信息时使用。会暂停任务执行，等待用户回复。')
    add_bullet(doc, 'question (必填): 向用户提出的问题，支持 Markdown 格式')

    # 5.12 task_done
    doc.add_heading('5.12 task_done — 标记完成', level=2)
    doc.add_paragraph('任务完成时调用，将摘要和代码路径写入 .agent_history.json。')
    add_bullet(doc, 'summary (必填): 任务完成总结')
    add_bullet(doc, 'code_path (可选): 涉及的主要代码路径')
    doc.add_paragraph('历史记录最多保留 20 条，超出时自动移除旧记录。')

    # 5.13 web_search
    doc.add_heading('5.13 web_search — 联网搜索', level=2)
    doc.add_paragraph('通过 DuckDuckGo Instant Answer API 进行联网搜索。需要先开启 enable_web_search 配置。'
                      '在 DeepSeek V4 Flash + Responses API 模式下，使用服务端原生搜索，效果更佳。')
    add_bullet(doc, 'query (必填): 搜索关键词或问题')

    # 5.14 open_file
    doc.add_heading('5.14 open_file — 打开文件', level=2)
    doc.add_paragraph('使用系统默认程序打开文件。支持图片、文档、网页等所有常见文件类型。')
    add_bullet(doc, 'path (必填): 文件路径')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 6. 配置与选项
    # ═══════════════════════════════════
    doc.add_heading('6. 配置与选项', level=1)

    doc.add_paragraph('所有配置项存储在 %LOCALAPPDATA%\\vibe_agent\\config.json 中。以下是完整的配置项说明：')

    # 基本配置
    doc.add_heading('6.1 基本配置', level=2)
    table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
    hdr = table.rows[0].cells
    hdr[0].text = '配置项'
    hdr[1].text = '默认值'
    hdr[2].text = '说明'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    configs = [
        ('language', 'zh_CN', '界面语言'),
        ('model', 'deepseek-v4-pro', '模型选择。可选 deepseek-v4-pro / deepseek-v4-flash'),
        ('api_base', 'https://api.deepseek.com', 'API 端点地址。支持自定义兼容端点'),
        ('project_root', '~/vibe_workspace', '默认工作区根目录'),
        ('max_steps', '128', '单次任务最大推理步数'),
        ('temperature', '1.0', '生成温度（0.0-2.0）'),
        ('think_level', '高', '推理深度：关 / 低 / 中 / 高'),
        ('max_tokens', '32767', '单次回复最大 token 数'),
        ('task_timeout', '0', '任务超时秒数。0 表示不限制'),
        ('enable_web_search', 'false', '是否启用联网搜索'),
        ('confirm_write_delete', 'true', '写/删文件前是否需要用户确认'),
        ('use_responses_api', 'true', 'Flash 模型是否使用 Responses API'),
    ]
    for key, default, desc in configs:
        row = table.add_row()
        cells = row.cells
        cells[0].text = key
        cells[1].text = default
        cells[2].text = desc

    # 记忆配置
    doc.add_heading('6.2 记忆系统配置', level=2)
    table2 = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '配置项'
    hdr2[1].text = '默认值'
    hdr2[2].text = '说明'
    for cell in hdr2:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    mem_configs = [
        ('memory', 'false', '是否启用记忆系统'),
        ('memory_mode', 'full', '记忆模式：full（完整）/ summary（摘要）'),
        ('max_rounds', '10', '完整模式下保留的最大对话轮数'),
    ]
    for key, default, desc in mem_configs:
        row = table2.add_row()
        cells = row.cells
        cells[0].text = key
        cells[1].text = default
        cells[2].text = desc

    # 插件配置
    doc.add_heading('6.3 插件系统配置', level=2)
    table3 = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '配置项'
    hdr3[1].text = '默认值'
    hdr3[2].text = '说明'
    for cell in hdr3:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    plugin_configs = [
        ('plugins_enabled', 'true', '是否启用插件系统'),
        ('plugin_dirs', '[]', '插件目录列表'),
        ('plugin_security_audit', 'warn', '安全审计级别：off / warn / block'),
        ('plugin_security_import_restrict', 'off', '导入限制：off / safe / strict'),
        ('plugin_security_require_permissions', 'false', '是否要求插件声明权限'),
        ('plugin_security_resource_limit', 'false', '是否启用插件资源限制'),
    ]
    for key, default, desc in plugin_configs:
        row = table3.add_row()
        cells = row.cells
        cells[0].text = key
        cells[1].text = default
        cells[2].text = desc

    # 异步架构配置
    doc.add_heading('6.4 异步架构配置', level=2)
    table4 = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '配置项'
    hdr4[1].text = '默认值'
    hdr4[2].text = '说明'
    for cell in hdr4:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    arch_configs = [
        ('sandbox_pool_max', '8', '沙箱池最大沙箱数'),
        ('sandbox_network_enabled', 'false', '沙箱是否允许网络访问'),
        ('file_io_queue_enabled', 'true', '是否启用文件 I/O 并发检测'),
        ('lifecycle_zombie_scan_seconds', '5', '僵尸进程扫描间隔（秒）'),
        ('resource_terminal_reserved_pct', '40', '终端保留资源百分比'),
    ]
    for key, default, desc in arch_configs:
        row = table4.add_row()
        cells = row.cells
        cells[0].text = key
        cells[1].text = default
        cells[2].text = desc

    doc.add_page_break()

    # ═══════════════════════════════════
    # 7. 插件系统
    # ═══════════════════════════════════
    doc.add_heading('7. 插件系统', level=1)

    doc.add_paragraph(
        'Vibe Coding Agent 内置了一个完整的插件框架，允许开发者扩展 Agent 的能力。'
        '插件可以注册自定义工具、监听生命周期钩子、访问会话上下文。'
    )

    doc.add_heading('7.1 内置官方插件', level=2)
    doc.add_paragraph('系统预装了以下官方插件：')
    official = [
        ('Code Reviewer (代码审查)', '对源代码文件执行全面的代码质量审查。检查文档字符串、异常处理、'
         '代码复杂度、命名规范、TODO/FIXME 标记、安全隐患等，并生成结构化的审查报告和评分。'),
        ('Dev Utilities (开发工具集)', '提供 UUID 生成、安全密码生成、哈希计算（MD5/SHA1/SHA256/SHA512）、'
         '时间戳与日期互转、项目代码行数统计等开发常用工具。'),
        ('Note Manager (笔记管理)', '在对话中创建、编辑、搜索结构化笔记。支持标签分类和全文搜索。'),
        ('Time Tracker (时间追踪)', '追踪 Agent 的任务执行时间，生成时间报告和效率分析。'),
    ]
    for name, desc in official:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        p.add_run(f'：{desc}')

    doc.add_heading('7.2 插件安全管理', level=2)
    doc.add_paragraph('插件加载前会自动经过多层安全检查：')
    checks = [
        'AST 源码审计 — 扫描危险模式（os.system、subprocess、eval、exec 等）',
        '导入限制 — 阻止插件导入 ctypes、subprocess、socket 等危险模块',
        '权限声明 — 可要求插件在 manifest.json 中声明所需权限',
        '资源限制 — 可限制插件的 CPU 时间（30s）和内存（512MB）',
        '安全等级 — off（关闭审计）/ warn（警告但允许）/ block（阻断加载）',
    ]
    for c in checks:
        add_bullet(doc, c)

    doc.add_paragraph('详细的插件开发指南请参见《NORP Vibe Coding Agent 插件开发说明》文档。')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 8. 安全机制
    # ═══════════════════════════════════
    doc.add_heading('8. 安全机制', level=1)

    doc.add_heading('8.1 路径边界', level=2)
    doc.add_paragraph(
        '所有文件操作均通过 _safe_path() 方法验证，确保路径限定在工作区根目录内。'
        '任何包含 ".." 或指向工作区外部的路径都会被拒绝。'
    )

    doc.add_heading('8.2 命令安全', level=2)
    doc.add_paragraph('Shell 命令执行前会检查危险模式：')
    dangerous = ['sudo', 'rm -rf /', 'mkfs', 'dd if=', '> /dev/sda', 'format c:']
    for d in dangerous:
        add_bullet(doc, d)

    doc.add_heading('8.3 操作确认', level=2)
    doc.add_paragraph(
        '开启 confirm_write_delete 选项后，write_file、delete_file、replace_in_file '
        '操作会弹出确认对话框，用户必须手动确认才会执行。'
    )

    doc.add_heading('8.4 API 密钥保护', level=2)
    doc.add_paragraph('API 密钥支持两种加密存储方式：')
    protections = [
        'win32crypt (Windows DPAPI) — 使用 Windows 数据保护 API 加密，仅当前用户可解密',
        'keyring — 使用系统密钥环服务存储',
    ]
    for p_text in protections:
        add_bullet(doc, p_text)

    doc.add_heading('8.5 沙箱隔离', level=2)
    doc.add_paragraph('支持两种沙箱模式：')
    sandboxes = [
        'Docker 容器 — 完全隔离的文件系统、网络、内存限制（推荐）',
        '子进程隔离 — Windows Job Object / Unix 进程组，无需 Docker',
    ]
    for s in sandboxes:
        add_bullet(doc, s)

    doc.add_heading('8.6 插件安全', level=2)
    doc.add_paragraph('详见第 7.2 节。多层安全机制确保第三方插件不会危害系统安全。')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 9. 异步架构
    # ═══════════════════════════════════
    doc.add_heading('9. 异步架构', level=1)

    doc.add_paragraph('Vibe Coding Agent 从最初的多线程架构重构为基于 asyncio 的异步架构，'
                      '解决了并发任务、资源竞争和僵尸进程三大核心问题。')

    doc.add_heading('9.1 沙箱池 (SandboxPool)', level=2)
    doc.add_paragraph('管理最多 8 个沙箱实例的异步池。任务需要沙箱时调用 acquire() 异步获取，'
                      '若无空闲则排队等待。使用完毕后 release() 归还。支持：')
    pool_features = [
        'Docker 容器 + 子进程隔离双模式',
        '路径映射（宿主路径 ↔ 沙箱内路径）',
        '命令执行结果路径反向映射',
        '进程组管理（Windows Job Object / Unix PGID）',
    ]
    for f in pool_features:
        add_bullet(doc, f)

    doc.add_heading('9.2 文件 I/O 队列 (FileIOQueue)', level=2)
    doc.add_paragraph('解决多任务并发读写同一文件时的冲突问题：')
    io_features = [
        '冲突检测：读-读无冲突，读-写/写-写冲突',
        'FIFO 排队：冲突操作自动排队等待',
        '防止写饥饿：后续读者在队列中有写者时也需等待',
        '30 秒超时：避免死锁',
    ]
    for f in io_features:
        add_bullet(doc, f)

    doc.add_heading('9.3 生命周期管理 (LifecycleManager)', level=2)
    doc.add_paragraph('解决用户停止任务后残留僵尸进程的问题：')
    lm_features = [
        '进程组绑定：每个任务的所有子进程注册到同一进程组',
        '级联终止：停止任务时杀整个进程组（而非单个进程）',
        '僵尸扫描：每 5 秒扫描已停止任务，确保进程全部终止',
        '用户等待保护：WAITING_USER 状态的任务不会被扫描器误杀',
    ]
    for f in lm_features:
        add_bullet(doc, f)

    doc.add_heading('9.4 权限级联 (PermissionCascade)', level=2)
    doc.add_paragraph('层级权限模型：系统 > 终端 > 插件根 > 插件子调用。'
                      '子操作权限 = 父级权限 ∩ 子级声明权限（取交集）。'
                      '任何操作不能超出其父级的权限范围。')

    doc.add_heading('9.5 资源隔离 (ResourceIsolator)', level=2)
    doc.add_paragraph('解决插件与终端抢资源的冲突：')
    ri_features = [
        '终端保留 40% 系统资源（优先级最高）',
        '插件池占 60%，按插件数均分',
        '每个插件独立配额：CPU 30s、内存 256MB、I/O 50MB',
        '配额耗尽时拒绝或排队',
    ]
    for f in ri_features:
        add_bullet(doc, f)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 10. 最佳实践
    # ═══════════════════════════════════
    doc.add_heading('10. 最佳实践', level=1)

    doc.add_heading('10.1 编写有效的任务指令', level=2)
    tips = [
        '明确具体：不要只说"做一个网站"，而是说"用 Flask + Bootstrap 做一个带用户登录的博客"',
        '提供上下文：如果有现有代码库，告诉 Agent 先去了解项目结构',
        '分步请求：复杂任务可以分成多个小步骤，逐步推进',
        '指定约束：明确技术栈、代码风格、文件结构等偏好',
    ]
    for tip in tips:
        add_bullet(doc, tip)

    doc.add_heading('10.2 Token 优化', level=2)
    token_tips = [
        '使用 replace_in_file 而非 write_file 修改大文件',
        '用 read_file 的行范围参数只读取需要的片段',
        '复杂任务分多次对话，避免单次对话上下文过长',
        '开启记忆系统（summary 模式）管理长期对话',
    ]
    for tip in token_tips:
        add_bullet(doc, tip)

    doc.add_heading('10.3 安全建议', level=2)
    safety_tips = [
        '保持 confirm_write_delete 开启，避免意外文件修改',
        '定期检查 plugin_dirs，移除不信任的插件目录',
        '使用 Docker 沙箱模式运行不信任的代码',
        '不要在公开场合分享 config.json 或 base.env 文件',
    ]
    for tip in safety_tips:
        add_bullet(doc, tip)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 11. 常见问题
    # ═══════════════════════════════════
    doc.add_heading('11. 常见问题 (FAQ)', level=1)

    faq = [
        ('Q: Agent 执行到一半卡住了怎么办？',
         'A: 点击「停止」按钮。系统会通过 LifecycleManager 杀掉整个进程组，不会留下僵尸进程。'),
        ('Q: 如何切换模型？',
         'A: 在设置中修改 model 字段。支持 deepseek-v4-pro（通用）和 deepseek-v4-flash（快速）。'
         'Flash 模型会自动启用 Responses API 获得更好的搜索体验。'),
        ('Q: 支持哪些 API 提供商？',
         'A: 默认支持 DeepSeek 官方端点。同时也支持任意兼容 OpenAI SDK 的自定义端点。'
         '联网搜索功能还兼容 DeepSeek 的 Anthropic 端点。'),
        ('Q: 插件加载失败怎么办？',
         'A: 检查设置中的 plugin_security_audit 级别。如果设为 "block"，任何有安全隐患的插件都会被拒绝加载。'
         '可以临时设为 "warn" 或 "off"，查看具体的审计报告后决定。'),
        ('Q: 文件操作被拒绝？',
         'A: 检查 confirm_write_delete 是否开启（默认开启）。如果开启了，写/删操作需要手动确认。'
         '另外确认目标路径是否在工作区根目录内。'),
        ('Q: 如何查看 Token 消耗？',
         'A: 右上角实时显示 input/output/tool_call 三类 token 的累计消耗。'
         '也可以在设置中点击「查询余额」获取 DeepSeek 账户的余额信息。'),
        ('Q: 多个会话会互相影响吗？',
         'A: 不会。每个会话拥有独立的事件队列、Agent 循环、对话历史和工作区。'
         '但文件 I/O 队列会检测跨会话的文件并发冲突并自动排队。'),
    ]
    for q, a in faq:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        doc.add_paragraph(a)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 12. 附录
    # ═══════════════════════════════════
    doc.add_heading('12. 附录', level=1)

    doc.add_heading('12.1 项目目录结构', level=2)
    add_code_block(doc, '''vibe_agent/
├── main.py                 # 程序入口
├── api.py                  # pywebview API 桥接
├── async_loop.py           # 异步 Agent 主循环
├── async_executor.py       # 异步工具执行器
├── config.py               # 配置管理
├── event_queue.py          # 事件队列
├── loop.py                 # 同步 Agent 循环（旧版）
├── executor.py             # 同步工具执行器（旧版）
├── sandbox_pool.py         # 沙箱池管理
├── file_io_queue.py        # 文件 I/O 队列
├── lifecycle_manager.py    # 生命周期管理
├── path_mapper.py          # 路径映射
├── permission_cascade.py   # 权限级联
├── resource_isolator.py    # 资源隔离
├── tools.py                # 内置工具定义
├── front.html              # 前端界面
├── plugin_system/          # 插件框架
│   ├── __init__.py
│   ├── manager.py          # 插件管理器
│   ├── security.py         # 安全审计
│   └── context.py          # 插件上下文
└── official_plugins/       # 官方插件
    ├── code_reviewer.py
    ├── dev_utilities.py
    ├── note_manager.py
    └── time_tracker.py''')

    doc.add_heading('12.2 API 调用模式', level=2)
    doc.add_paragraph('Agent 支持三种 API 调用模式，根据 model 和 base_url 自动切换：')

    table5 = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '模式'
    hdr5[1].text = '条件'
    hdr5[2].text = '特点'
    for cell in hdr5:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    api_modes = [
        ('Responses API', 'Flash 模型 + 官方端点', '服务端原生搜索、语义化事件流、无状态调用'),
        ('Chat Completions', 'Pro 模型 或 自定义端点', '标准 OpenAI 格式、工具调用、流式推理'),
        ('Anthropic Messages', '官方端点 + 启用搜索', 'Anthropic 原生搜索工具、独立系统提示'),
    ]
    for mode, cond, feat in api_modes:
        row = table5.add_row()
        cells = row.cells
        cells[0].text = mode
        cells[1].text = cond
        cells[2].text = feat

    doc.add_heading('12.3 版本信息', level=2)
    doc.add_paragraph('NORP Vibe Coding Agent v1.0')
    doc.add_paragraph(f'Copyright © 2026 xingluosama')
    doc.add_paragraph(f'文档生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')

    # ── 保存 ──
    output_path = os.path.join(OUTPUT_DIR, '使用文档.docx')
    doc.save(output_path)
    print(f'[OK] User manual generated: {output_path}')
    return output_path


def create_plugin_dev_guide():
    """创建插件开发说明 .docx。"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '微软雅黑'
        hs.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    # ── 封面 ──
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('NORP Vibe Coding Agent')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('插件开发说明')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'版本 1.0 | {datetime.date.today().strftime("%Y年%m月")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 1. 概述
    # ═══════════════════════════════════
    doc.add_heading('1. 插件系统概述', level=1)

    doc.add_paragraph(
        'NORP Vibe Coding Agent 的插件系统允许开发者通过标准化的接口扩展 Agent 的能力。'
        '插件可以注册自定义工具（供 AI 调用）、监听 15 个生命周期钩子（干预 Agent 行为）、'
        '以及访问会话上下文获取环境信息。'
    )

    doc.add_heading('1.1 插件能做什么', level=2)
    capabilities = [
        '注册自定义工具：Agent 可以像使用内置工具一样调用你的工具函数',
        '拦截生命周期：在任务开始/结束、工具调用前后、推理流式输出等时刻执行自定义逻辑',
        '修改数据流：before_step、before_tool_call、after_tool_call 钩子可以修改传入的数据',
        '访问上下文：通过 PluginContext 获取工作区路径、配置快照、Token 用量等信息',
        '持久化状态：通过 context.storage 在单次任务中跨钩子保持状态',
    ]
    for c in capabilities:
        add_bullet(doc, c)

    doc.add_heading('1.2 插件不能做什么', level=2)
    limits = [
        '不能执行任意系统命令（受安全审计和导入限制约束）',
        '不能访问工作区外的文件',
        '不能导入 ctypes、subprocess 等危险模块（受导入限制策略约束）',
        '不能修改其他插件的数据',
        '单个钩子执行超时 5 秒（HOOK_TIMEOUT）',
    ]
    for l in limits:
        add_bullet(doc, l)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 2. 插件文件结构
    # ═══════════════════════════════════
    doc.add_heading('2. 插件文件结构', level=1)

    doc.add_paragraph('插件支持两种文件组织形式：')

    doc.add_heading('2.1 单文件插件（推荐用于简单插件）', level=2)
    doc.add_paragraph('直接在插件目录下放置一个 .py 文件：')
    add_code_block(doc, '''plugin_dirs/
  my_plugin.py        ← 一个文件就是一个插件''')
    doc.add_paragraph('文件名（不含扩展名）即为插件名称。')

    doc.add_heading('2.2 包格式插件（推荐用于复杂插件）', level=2)
    doc.add_paragraph('包含 manifest.json 的目录：')
    add_code_block(doc, '''plugin_dirs/
  my_plugin/
    manifest.json     ← 插件元数据（必须）
    plugin.py         ← 入口文件（默认，可在 manifest 中指定）
    utils.py          ← 其他辅助模块
    data/             ← 数据文件''')

    doc.add_heading('2.3 manifest.json 格式', level=2)
    add_code_block(doc, '''{
    "name": "My Plugin",
    "version": "1.0.0",
    "publisher": "Your Name",
    "description": "插件功能说明",
    "entry": "plugin.py",
    "enabled": true,
    "permissions": ["file_read", "network"]
}''')
    doc.add_paragraph('说明：')
    manifest_fields = [
        'name: 插件显示名称（可选，未提供则用目录名）',
        'version: 语义版本号',
        'publisher / author: 作者名',
        'description: 功能描述',
        'entry: 入口文件名（默认 plugin.py）',
        'enabled: 是否启用',
        'permissions: 权限声明列表（需开启 require_permissions 配置）',
    ]
    for f in manifest_fields:
        add_bullet(doc, f)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 3. 插件元数据
    # ═══════════════════════════════════
    doc.add_heading('3. 插件元数据声明', level=1)

    doc.add_paragraph('每个插件必须（或建议）在模块顶层声明以下常量：')

    add_code_block(doc, '''# 必须声明
PLUGIN_NAME = "My Plugin"

# 建议声明
PLUGIN_PUBLISHER = "Author Name"
PLUGIN_VERSION = "1.0.0"
PLUGIN_DESCRIPTION = "描述插件功能"''')

    doc.add_paragraph('如果使用了 manifest.json，其中的 name/version/publisher/description '
                      '优先级高于模块常量的声明。')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 4. 工具注册
    # ═══════════════════════════════════
    doc.add_heading('4. 工具注册', level=1)

    doc.add_heading('4.1 TOOLS 常量', level=2)
    doc.add_paragraph('在模块中定义 TOOLS 列表，每个元素为 OpenAI Function Schema 字典：')

    add_code_block(doc, '''TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "my_tool",
            "description": "工具的功能说明（AI 据此判断何时调用）",
            "parameters": {
                "type": "object",
                "properties": {
                    "param1": {
                        "type": "string",
                        "description": "参数1的说明"
                    },
                    "param2": {
                        "type": "integer",
                        "description": "参数2的说明",
                        "default": 42
                    }
                },
                "required": ["param1"],
                "additionalProperties": False
            }
        }
    }
]''')

    doc.add_heading('4.2 工具命名规范', level=2)
    naming_tips = [
        '使用 snake_case 命名风格（小写 + 下划线）',
        '工具名应在所有插件中唯一。如果重复，后加载的插件会报错',
        'description 字段至关重要——AI 完全依赖它来判断是否应该调用你的工具',
        '参数描述应清晰、具体。提供 default 值可以简化 AI 的调用',
    ]
    for tip in naming_tips:
        add_bullet(doc, tip)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 5. 工具执行函数
    # ═══════════════════════════════════
    doc.add_heading('5. 工具执行函数', level=1)

    doc.add_heading('5.1 execute() 函数签名', level=2)
    add_code_block(doc, '''def execute(tool_name: str, args: dict, context: PluginContext) -> str:
    """工具执行入口。

    Args:
        tool_name: 被调用的工具名称
        args: 工具参数字典（由 AI 填充）
        context: 插件上下文对象（只读环境信息 + 可写 storage）

    Returns:
        str: 工具执行结果（将作为 tool message 返回给 AI）
    """
    if tool_name == "my_tool":
        return _handle_my_tool(args, context)
    return f"Unknown tool: {tool_name}"''')

    doc.add_heading('5.2 路径安全', level=2)
    doc.add_paragraph('如果你的工具操作文件，必须确保路径在工作区范围内：')
    add_code_block(doc, '''def _handle_my_tool(args, context):
    file_path = args.get("file_path", "")
    # 安全的路径拼接
    full_path = os.path.join(context.project_root, file_path)
    full_path = os.path.normpath(full_path)

    # 验证路径在工作区内
    if not full_path.startswith(context.project_root):
        return "Error: path out of bounds"

    with open(full_path, "r", encoding="utf-8") as f:
        return f.read()''')

    doc.add_heading('5.3 日志记录', level=2)
    doc.add_paragraph('使用 context.logger 记录日志：')
    add_code_block(doc, '''context.logger.info("Starting code review...")
context.logger.warn("File is unusually large")
context.logger.error("Failed to parse file")
context.logger.debug("Processing line 42")''')
    doc.add_paragraph('日志会输出到控制台并写入 %LOCALAPPDATA%\\vibe_agent\\plugin.log。')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 6. 钩子系统
    # ═══════════════════════════════════
    doc.add_heading('6. 钩子系统', level=1)

    doc.add_paragraph('插件系统提供 15 个生命周期钩子，分 4 个层级。只需在模块中定义同名函数即可接收事件。')

    doc.add_heading('6.1 L1 — 生命周期钩子', level=2)

    doc.add_heading('on_agent_init(context)', level=3)
    doc.add_paragraph('Agent 初始化完成时触发。适合初始化插件状态、设置计数器等。')
    add_code_block(doc, '''def on_agent_init(context):
    context.storage["call_count"] = 0
    context.logger.info("Plugin loaded!")''')

    doc.add_heading('on_agent_shutdown(context)', level=3)
    doc.add_paragraph('Agent 关闭时触发。适合输出统计信息、清理资源。')

    doc.add_heading('6.2 L2 — 任务钩子', level=2)

    doc.add_heading('on_task_start(task_text: str, context)', level=3)
    doc.add_paragraph('每次用户发送新任务时触发。task_text 是用户的原始输入。')
    doc.add_paragraph('用途：检测任务关键词，自动启用相关功能。')

    doc.add_heading('on_task_done(summary: str, final_reply: str, context)', level=3)
    doc.add_paragraph('任务成功完成时触发。summary 是 task_done 的摘要，final_reply 是最终回复。')

    doc.add_heading('on_task_error(error_msg: str, context)', level=3)
    doc.add_paragraph('任务执行异常时触发。')

    doc.add_heading('on_task_stopped(context)', level=3)
    doc.add_paragraph('用户手动停止任务时触发。')

    doc.add_heading('on_task_timeout(elapsed: float, context)', level=3)
    doc.add_paragraph('任务超时时触发。elapsed 是已用秒数。')

    doc.add_heading('6.3 L3 — 步骤钩子', level=2)

    doc.add_heading('before_step(step: int, messages: list, context) -> list | None', level=3)
    doc.add_paragraph('每轮推理步骤之前触发。可以修改 messages 列表（如注入额外上下文）。'
                      '返回修改后的 messages 列表，或返回 None 表示不修改。')

    doc.add_heading('after_step(step: int, reasoning: str, content: str, tool_calls: list, context)', level=3)
    doc.add_paragraph('每轮推理步骤之后触发。可获取该步骤的推理过程、输出内容和工具调用。')

    doc.add_heading('before_tool_call(tool_name: str, args: dict, context) -> dict | None', level=3)
    doc.add_paragraph('工具执行前触发。可以修改 args 或返回 None 阻止执行。'
                      '这是最常用的钩子之一，适合做参数校验、日志记录、调用计数。')

    doc.add_heading('after_tool_call(tool_name: str, args: dict, result: str, context) -> str', level=3)
    doc.add_paragraph('工具执行后触发。可以修改或增强工具返回结果。')

    doc.add_heading('on_user_input_required(question: str, context)', level=3)
    doc.add_paragraph('Agent 调用 ask_user 等待用户输入时触发。')

    doc.add_heading('6.4 L4 — 流式事件钩子', level=2)

    doc.add_heading('on_reasoning(token: str, context)', level=3)
    doc.add_paragraph('AI 输出推理过程时逐 token 触发。用于实时展示或分析推理内容。')

    doc.add_heading('on_content(token: str, context)', level=3)
    doc.add_paragraph('AI 输出回复内容时逐 token 触发。')

    doc.add_heading('on_event(event_type: str, data: str, context)', level=3)
    doc.add_paragraph('通用事件钩子，接收所有事件类型。')

    doc.add_heading('on_usage_update(usage: dict, context)', level=3)
    doc.add_paragraph('Token 用量更新时触发。usage 包含 input_tokens、output_tokens、tool_call_tokens。')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 7. PluginContext API
    # ═══════════════════════════════════
    doc.add_heading('7. PluginContext API 参考', level=1)

    doc.add_paragraph('PluginContext 是传递给所有钩子和 execute() 的上下文对象。')

    doc.add_heading('7.1 属性', level=2)
    ctx_attrs = [
        ('context.project_root', 'str', '当前工作区根目录的绝对路径'),
        ('context.app_dir', 'str', '应用数据目录（%LOCALAPPDATA%\\vibe_agent）'),
        ('context.config', 'dict', '当前 config.json 的只读快照'),
        ('context.storage', 'dict', '插件私有存储字典。在单次 Agent 生命周期内跨钩子持久化'),
        ('context.logger', 'SimpleLogger', '插件专用日志记录器'),
        ('context.current_step', 'int', '当前 ReAct 步骤编号'),
        ('context.total_usage', 'dict', '累计 Token 用量快照'),
    ]
    for name, typ, desc in ctx_attrs:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.add_run(f' ({typ})：{desc}')

    doc.add_heading('7.2 SimpleLogger API', level=2)
    logger_methods = [
        'logger.info(msg) — 信息日志',
        'logger.warn(msg) — 警告日志',
        'logger.error(msg) — 错误日志',
        'logger.debug(msg) — 调试日志',
    ]
    for m in logger_methods:
        add_bullet(doc, m)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 8. 安全机制
    # ═══════════════════════════════════
    doc.add_heading('8. 插件安全机制详解', level=1)

    doc.add_heading('8.1 源码审计', level=2)
    doc.add_paragraph('插件加载前，系统会对其源码进行 AST 级别的静态分析，检测以下危险模式：')

    doc.add_heading('严重级 (CRITICAL) — 阻断加载', level=3)
    critical_items = [
        'Shell 执行：os.system()、os.popen()、subprocess.call()、subprocess.run() 等',
        '代码执行：eval()、exec()、compile()、__import__()',
        '原生代码：ctypes、cffi',
        '进程终止：os._exit()、os.kill()',
        '导入绕过：importlib.import_module()',
    ]
    for item in critical_items:
        add_bullet(doc, item)

    doc.add_heading('警告级 (WARNING) — 允许但记录', level=3)
    warning_items = [
        '文件操作：os.remove()、shutil.rmtree()、shutil.move()',
        '网络访问：socket、http、urllib、requests、ftplib 等',
        '反序列化：pickle、marshal、yaml.unsafe_load()',
        '系统操纵：sys.modules、sys.exit()、builtins 修改',
    ]
    for item in warning_items:
        add_bullet(doc, item)

    doc.add_heading('8.2 导入限制', level=2)
    doc.add_paragraph('三种导入限制策略（通过 plugin_security_import_restrict 配置）：')
    import_policies = [
        ('off（默认）', '不限制。插件可导入任意模块。'),
        ('safe', '阻止导入危险模块（subprocess、ctypes、socket、pickle 等）。仅对插件代码生效。'),
        ('strict', '严格白名单模式。仅允许导入 STRICT_SAFE_MODULES 中的模块（json、re、datetime 等基础库）。'
         '其他所有模块均被阻止。'),
    ]
    for name, desc in import_policies:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        p.add_run(f'：{desc}')

    doc.add_heading('8.3 权限声明', level=2)
    doc.add_paragraph('当开启 require_permissions 配置时，插件必须在 manifest.json 中声明所需权限：')
    add_code_block(doc, '''{
    "permissions": ["file_read", "file_write", "process", "network"]
}''')
    doc.add_paragraph('系统会将安全审计发现的风险操作与声明的权限进行比对，缺失的权限会导致加载失败。')

    doc.add_heading('8.4 资源限制', level=2)
    doc.add_paragraph('当开启 resource_limit 配置时，插件加载过程中会启用资源限制：')
    resource_limits_list = [
        'CPU 时间限制：30 秒',
        '内存限制：512 MB',
        '基于 Unix signal.SIGALRM 的 CPU 超时',
        '基于 resource.setrlimit 的内存限制（Unix）',
    ]
    for r in resource_limits_list:
        add_bullet(doc, r)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 9. 完整示例
    # ═══════════════════════════════════
    doc.add_heading('9. 完整插件示例', level=1)

    doc.add_paragraph('以下是一个完整的插件示例，展示了工具注册、执行和钩子的用法：')

    add_code_block(doc, '''# ── 元数据 ──
PLUGIN_NAME = "Hello World Plugin"
PLUGIN_PUBLISHER = "Your Name"
PLUGIN_VERSION = "1.0.0"
PLUGIN_DESCRIPTION = "一个演示插件"

import os

# ── 工具注册 ──
TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "hello_world",
            "description": "向指定的人问好",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {
                        "type": "string",
                        "description": "要问候的名字"
                    }
                },
                "required": ["name"],
                "additionalProperties": False
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "count_files",
            "description": "统计工作区中的文件数量",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
                "additionalProperties": False
            }
        }
    }
]

# ── 工具执行 ──
def execute(tool_name: str, args: dict, context) -> str:
    if tool_name == "hello_world":
        name = args.get("name", "World")
        context.storage["greet_count"] = \\
            context.storage.get("greet_count", 0) + 1
        return f"Hello, {name}! 👋"

    if tool_name == "count_files":
        count = 0
        for root, dirs, files in os.walk(context.project_root):
            # 跳过隐藏目录
            dirs[:] = [d for d in dirs if not d.startswith(".")]
            count += len(files)
        return f"工作区共有 {count} 个文件。"

    return f"Unknown tool: {tool_name}"

# ── 钩子 ──
def on_agent_init(context):
    context.storage["greet_count"] = 0
    context.logger.info("Hello World plugin loaded!")

def on_agent_shutdown(context):
    count = context.storage.get("greet_count", 0)
    context.logger.info(f"Said hello {count} time(s)")

def on_task_start(task_text: str, context):
    if "hello" in task_text.lower():
        context.logger.info("Hello-related task detected!")

def after_tool_call(tool_name: str, args: dict, result: str, context):
    if tool_name == "hello_world":
        # 给回复添加时间戳
        from datetime import datetime
        ts = datetime.now().strftime("%H:%M:%S")
        return f"{result}\\n（执行时间: {ts})"
    return result''')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 10. 调试与测试
    # ═══════════════════════════════════
    doc.add_heading('10. 调试与测试', level=1)

    doc.add_heading('10.1 查看审计报告', level=2)
    doc.add_paragraph('插件加载后，可以在前端界面查看每个插件的安全审计结果。'
                      '也可以通过 API get_plugin_audit_results() 获取详细信息。')

    doc.add_heading('10.2 日志调试', level=2)
    doc.add_paragraph('利用 context.logger 输出调试信息。日志文件位于：')
    add_code_block(doc, '%LOCALAPPDATA%\\vibe_agent\\plugin.log')
    doc.add_paragraph('也可以在控制台查看实时日志输出（如果以命令行方式启动）。')

    doc.add_heading('10.3 常用调试技巧', level=2)
    debug_tips = [
        '先在 audit=warn 模式下测试插件，确认没有严重安全问题',
        '使用 before_tool_call 钩子打印工具调用参数，了解 Agent 如何调用你的工具',
        '使用 after_step 钩子查看每轮推理的完整上下文',
        '在 context.storage 中记录调用次数、参数等，便于在 on_agent_shutdown 中输出汇总',
    ]
    for tip in debug_tips:
        add_bullet(doc, tip)

    doc.add_heading('10.4 常见问题', level=2)

    faq_dev = [
        ('Q: 插件加载后没有出现在插件列表中？',
         'A: 检查以下几点：① 插件文件是否在 plugin_dirs 配置的目录中；② 是否定义了 PLUGIN_NAME 常量；'
         '③ 是否定义了 TOOLS 或至少一个钩子函数；④ 安全审计是否通过了（检查 audit 级别）。'),
        ('Q: 工具未被 AI 调用？',
         'A: 检查 TOOLS 中 function.description 是否足够清晰描述工具用途。AI 完全依赖描述来决定是否调用。'
         '如果描述模糊或不准确，AI 可能不会调用你的工具。'),
        ('Q: 钩子函数没有被触发？',
         'A: 确认函数名与 HOOK_NAMES 列表中的名称完全一致（区分大小写）。'
         '确认插件已成功加载（在插件列表中显示为 enabled）。'),
        ('Q: 如何在钩子中阻止操作？',
         'A: before_tool_call 钩子返回 None 可阻止工具执行。before_step 钩子返回修改后的 messages 列表。'
         '注意：只有标注为 _MUTATING_HOOKS 的钩子支持修改数据流。'),
    ]
    for q, a in faq_dev:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        doc.add_paragraph(a)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 11. 附录：参考实现
    # ═══════════════════════════════════
    doc.add_heading('11. 附录：官方插件参考', level=1)

    doc.add_paragraph('建议参考以下官方插件的源码来学习最佳实践：')

    ref_plugins = [
        ('code_reviewer.py', 'Code Reviewer', '展示了如何注册复杂工具（多参数、严格度选项）、'
         '如何利用钩子检测任务关键词、如何维护跨钩子状态。'),
        ('dev_utilities.py', 'Dev Utilities', '展示了如何注册多个工具、工具的模块化实现模式、'
         '安全性（secrets 模块生成密码）、结果格式化。'),
        ('note_manager.py', 'Note Manager', '展示了如何实现有状态插件（创建/编辑/搜索笔记）、'
         '如何使用 context.storage 维护数据。'),
        ('time_tracker.py', 'Time Tracker', '展示了如何利用多个钩子（on_task_start/on_task_done）'
         '追踪任务执行时间、生成效率报告。'),
    ]
    for filename, name, desc in ref_plugins:
        p = doc.add_paragraph()
        run = p.add_run(f'{filename} ({name})')
        run.bold = True
        doc.add_paragraph(desc)

    doc.add_heading('11.1 钩子函数完整列表', level=2)
    add_code_block(doc, '''HOOK_NAMES = [
    # L1 – Lifecycle
    "on_agent_init",
    "on_agent_shutdown",
    # L2 – Task
    "on_task_start",
    "on_task_done",
    "on_task_error",
    "on_task_stopped",
    "on_task_timeout",
    # L3 – Step
    "before_step",       # 可变（可返回修改后的 messages）
    "after_step",
    "before_tool_call",  # 可变（可返回修改后的 args 或 None 阻止）
    "after_tool_call",   # 可变（可返回修改后的 result）
    "on_user_input_required",
    # L4 – Streaming events
    "on_reasoning",
    "on_content",
    "on_event",
    "on_usage_update",
]''')

    doc.add_heading('11.2 安全审计分类完整列表', level=2)
    doc.add_paragraph('以下为 AST 审计器检测的完整危险模式分类：')
    add_code_block(doc, '''- shell_exec:        os.system, os.popen, subprocess.*
- code_exec:         eval, exec, compile, __import__
- import_bypass:     importlib.import_module
- native_exec:       ctypes, cffi
- process_terminate: os._exit, os.kill
- file_delete:       os.remove, os.unlink, shutil.rmtree, os.rmdir
- file_move:         shutil.move
- network:           socket, http, urllib, requests, ftplib, smtplib, poplib, telnetlib
- deserialization:   pickle, marshal, yaml.unsafe_load
- sys_manipulation:  sys.modules, sys.setprofile, sys.settrace
- builtin_override:  modifying builtins
- missing_permissions: 未声明所需权限''')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('— 文档结束 —')
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'Copyright © 2026 xingluosama | {datetime.date.today().strftime("%Y年%m月%d日")}')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── 保存 ──
    output_path = os.path.join(OUTPUT_DIR, '插件开发说明.docx')
    doc.save(output_path)
    print(f'[OK] Plugin dev guide generated: {output_path}')
    return output_path


if __name__ == '__main__':
    print('Generating documents...')
    path1 = create_user_manual()
    path2 = create_plugin_dev_guide()
    print(f'\nUsage manual: {path1}')
    print(f'Plugin dev guide: {path2}')
    print('Done!')
