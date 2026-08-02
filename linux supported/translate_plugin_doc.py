"""
翻译 插件开发说明.docx 为英语，生成 plugin_dev_guide_en.docx。
保留原文档格式（段落样式、表格结构、run级别格式）。
使用 MyMemory 翻译 API，长文本自动分块翻译。
"""
import time
import sys
import io
import docx
from docx import Document
from deep_translator import MyMemoryTranslator

# Fix Windows console encoding for emoji/Chinese
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')


def translate_en(text: str) -> str:
    """中文 → 英语，长文本自动分块"""
    if not text.strip():
        return text

    # Split long text into chunks of ~450 chars
    max_chunk = 450
    if len(text) <= max_chunk:
        return _translate_chunk(text)

    # Split by lines to keep code blocks together
    lines = text.split('\n')
    chunks = []
    current = ""
    for line in lines:
        if len(current) + len(line) + 1 > max_chunk and current:
            chunks.append(current)
            current = line
        else:
            current = current + '\n' + line if current else line
    if current:
        chunks.append(current)

    # Translate each chunk
    results = []
    for chunk in chunks:
        result = _translate_chunk(chunk)
        results.append(result)
        time.sleep(0.15)

    return '\n'.join(results)


def _is_valid_result(text: str) -> bool:
    """Check if translation result looks valid (not an API error message)"""
    if not text:
        return False
    # MyMemory error messages
    invalid_patterns = [
        'INVALID SOURCE LANGUAGE',
        'INVALID TARGET LANGUAGE',
        'NO CONTENT',
        'TEXT LENGTH NEED TO BE BETWEEN',
        'NO QUOTA',
        'QUOTA EXCEEDED',
        'INVALID REQUEST',
    ]
    upper = text.upper()
    for p in invalid_patterns:
        if p in upper:
            return False
    return True


def _translate_chunk(text: str) -> str:
    """翻译单个文本块"""
    if not text.strip():
        return text
    # If text is mostly code/English, don't translate
    if _is_mostly_code(text):
        return text
    try:
        result = MyMemoryTranslator(source='zh-CN', target='en-GB').translate(text)
        if result and result != text and _is_valid_result(result):
            return result
    except Exception:
        pass
    try:
        result = MyMemoryTranslator(source='auto', target='en-GB').translate(text)
        if result and result != text and _is_valid_result(result):
            return result
    except Exception:
        pass
    return text


def _is_mostly_code(text: str) -> bool:
    """判断是否主要是代码（不需要翻译）"""
    code_indicators = ['def ', 'import ', 'return ', '{', '}', 'PLUGIN_',
                       'TOOLS', 'HOOK_NAMES', 'class ', 'if __name__',
                       'except ', 'try:', 'from ', 'self.', '  ', '\t']
    # If it looks like Python code
    lines = text.strip().split('\n')
    code_lines = sum(1 for l in lines if any(ind in l for ind in code_indicators))
    if len(lines) > 0 and code_lines / len(lines) > 0.3:
        return True
    return False


def copy_paragraph_format(src_para, dst_para):
    dst_para.style = src_para.style
    dst_para.alignment = src_para.alignment


def copy_run_format(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb


def translate_document(src_path: str, dst_path: str):
    print(f"Translating: {src_path} -> {dst_path}")

    src_doc = Document(src_path)
    dst_doc = Document()

    # 复制页面设置
    for i, section in enumerate(src_doc.sections):
        if i >= len(dst_doc.sections):
            dst_doc.add_section()
        dst_section = dst_doc.sections[i]
        dst_section.page_width = section.page_width
        dst_section.page_height = section.page_height
        dst_section.left_margin = section.left_margin
        dst_section.right_margin = section.right_margin
        dst_section.top_margin = section.top_margin
        dst_section.bottom_margin = section.bottom_margin

    # 处理段落
    total_paras = len(src_doc.paragraphs)
    for idx, para in enumerate(src_doc.paragraphs):
        if idx % 10 == 0:
            print(f"  Paragraphs: {idx}/{total_paras}")

        text = para.text.strip()
        if not text and not para.runs:
            dst_doc.add_paragraph()
            continue

        # 翻译（自动跳过大段代码）
        translated = translate_en(text) if text else ""

        new_para = dst_doc.add_paragraph()
        copy_paragraph_format(para, new_para)

        if para.runs:
            first_run = para.runs[0]
            new_run = new_para.add_run(translated)
            copy_run_format(first_run, new_run)
        else:
            new_para.add_run(translated)

        if idx % 5 == 0 and idx > 0:
            time.sleep(0.2)

    # 处理表格
    print(f"  Processing {len(src_doc.tables)} tables...")
    for ti, table in enumerate(src_doc.tables):
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        new_table = dst_doc.add_table(rows=num_rows, cols=num_cols)

        if table.style:
            new_table.style = table.style

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                src_cell_text = cell.text
                translated = translate_en(src_cell_text) if src_cell_text.strip() else src_cell_text
                new_table.rows[ri].cells[ci].text = translated
                if cell.paragraphs:
                    src_cell_para = cell.paragraphs[0]
                    dst_cell_para = new_table.rows[ri].cells[ci].paragraphs[0]
                    copy_paragraph_format(src_cell_para, dst_cell_para)
                time.sleep(0.1)

    dst_doc.save(dst_path)
    print(f"  Saved: {dst_path}")


def main():
    translate_document("插件开发说明.docx", "plugin_dev_guide_en.docx")
    print("\nDone! Output: plugin_dev_guide_en.docx")


if __name__ == "__main__":
    main()
