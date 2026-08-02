"""
Generate plugin_develop_guide.docx — NORP Vibe Coding Agent Plugin Development Guide.
Uses python-docx to create a professionally formatted Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells_text, bold=False, header=False):
    """Add a row to a table with formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Consolas'
        if bold:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '2B579A')
    return row


def add_code_block(doc, code_text):
    """Add a code block with monospace font and light background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add small spacing after code block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)


def create_document():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('NORP Vibe Coding Agent')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Plugin Development Guide')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_p.add_run('Version 1.0 — August 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    doc.add_paragraph()

    copyright_p = doc.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_p.add_run('Copyright \u00a9 2026 xingluosama. All rights reserved.')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()

    # ── Table of Contents ──
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Overview',
        '2. Plugin Structure',
        '   2.1 Single-File Plugin',
        '   2.2 Package Plugin with manifest.json',
        '3. Plugin Metadata',
        '4. Tool Registration',
        '   4.1 OpenAI Function Schema Format',
        '   4.2 Complete TOOLS Example',
        '5. The execute() Function',
        '   5.1 Signature',
        '   5.2 Dispatching Pattern',
        '   5.3 Return Value',
        '6. Lifecycle Hooks',
        '   6.1 Hook Categories',
        '   6.2 L1 \u2014 Lifecycle Hooks',
        '   6.3 L2 \u2014 Task Hooks',
        '   6.4 L3 \u2014 Step Hooks (Mutating)',
        '   6.5 L4 \u2014 Streaming Event Hooks',
        '   6.6 Mutating vs Non-Mutating Hooks',
        '7. PluginContext API',
        '   7.1 Attributes',
        '   7.2 SimpleLogger',
        '8. Security System',
        '   8.1 AST Source Audit',
        '   8.2 Import Restriction',
        '   8.3 Permission Declaration',
        '   8.4 Resource Limits',
        '   8.5 Security Configuration Summary',
        '9. manifest.json Reference',
        '10. Step-by-Step Tutorial',
        '   10.1 Create the File',
        '   10.2 Declare Metadata',
        '   10.3 Register a Tool',
        '   10.4 Implement execute()',
        '   10.5 Add Lifecycle Hooks',
        '   10.6 Test Your Plugin',
        '11. Best Practices',
        '12. Troubleshooting',
        '13. Appendix',
        '   13.1 Complete Hook Reference',
        '   13.2 Dangerous Pattern Registry',
        '   13.3 PluginContext Quick Reference',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 1: OVERVIEW
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('1. Overview', level=1)

    doc.add_paragraph(
        'The NORP Vibe Coding Agent plugin system allows developers to extend the Agent\'s capabilities '
        'by registering custom tools, listening to lifecycle hooks, and accessing session context. '
        'Plugins are loaded at application startup from configurable plugin directories and run '
        'within a security-audited environment.'
    )

    doc.add_heading('What Plugins Can Do', level=2)
    capabilities = [
        'Register custom OpenAI function-schema tools that the AI Agent can invoke.',
        'Listen to 15 lifecycle hooks across 4 layers (lifecycle, task, step, streaming).',
        'Read read-only session context (project root, config snapshot, token usage).',
        'Persist arbitrary data in per-plugin storage across hooks within a session.',
        'Log messages via a built-in per-plugin logger.',
        'Mutate data flow through designated hooks (before_step, before_tool_call, after_tool_call).',
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_heading('What Plugins Cannot Do', level=2)
    limits = [
        'Execute arbitrary shell commands (blocked by security audit).',
        'Import dangerous modules such as subprocess, ctypes, or socket (blocked by import restrictor).',
        'Access files outside the workspace (path boundary enforcement).',
        'Modify system-level configuration or other plugins\' storage.',
        'Crash the Agent \u2014 all hook exceptions are silently caught.',
    ]
    for lim in limits:
        doc.add_paragraph(lim, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 2: PLUGIN STRUCTURE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('2. Plugin Structure', level=1)

    doc.add_paragraph(
        'Plugins can be placed in any directory listed under plugin_dirs in config.json. '
        'The PluginManager discovers plugins by scanning these directories and supports '
        'two layout styles:'
    )

    # 2.1 Single-File
    doc.add_heading('2.1 Single-File Plugin', level=2)
    doc.add_paragraph(
        'The simplest form. A single .py file placed directly in a plugin directory. '
        'The file name (without .py) becomes the default plugin name.'
    )

    doc.add_paragraph('Directory layout:')
    add_code_block(doc, 'plugins/\n  my_tool.py          # single-file plugin')

    doc.add_paragraph(
        'Notes:\n'
        '\u2022 The file must not be named __init__.py (these are skipped).\n'
        '\u2022 If the file defines PLUGIN_NAME, that value takes priority over the file name.'
    )

    # 2.2 Package Plugin
    doc.add_heading('2.2 Package Plugin with manifest.json', level=2)
    doc.add_paragraph(
        'For more complex plugins, use a dedicated subdirectory containing a manifest.json '
        'and an entry-point .py file.'
    )

    doc.add_paragraph('Directory layout:')
    add_code_block(doc, 'plugins/\n  fancy_tool/\n    manifest.json       # metadata (required for detection)\n    plugin.py            # entry point (default, can be customized)')

    doc.add_paragraph(
        'The manifest.json file is required for the PluginManager to recognize the directory '
        'as a plugin package. Without it, the directory is silently skipped.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 3: PLUGIN METADATA
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('3. Plugin Metadata', level=1)

    doc.add_paragraph(
        'Every plugin must declare its identity through module-level constants. '
        'These are read by PluginManager._load_from_file() after the module is imported.'
    )

    doc.add_heading('Required Constants', level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0]
    for i, text in enumerate(['Constant', 'Type', 'Default', 'Description']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    rows_data = [
        ('PLUGIN_NAME', 'str', '(required)', 'Human-readable plugin name shown in the UI.'),
        ('PLUGIN_PUBLISHER', 'str', '(required)', 'Author or organization name.'),
        ('PLUGIN_VERSION', 'str', '"0.0.0"', 'Semantic version string. manifest.json takes priority if both present.'),
        ('PLUGIN_DESCRIPTION', 'str', '""', 'Short description shown in the plugin list UI.'),
    ]
    for row_data in rows_data:
        add_table_row(table, row_data)

    doc.add_paragraph()
    doc.add_paragraph('Example:')
    add_code_block(doc, 'PLUGIN_NAME = "My Awesome Tool"\nPLUGIN_PUBLISHER = "Your Name"\nPLUGIN_VERSION = "1.0.0"\nPLUGIN_DESCRIPTION = "Does something amazing with AI agent workflows."')

    doc.add_paragraph(
        'Important: If manifest.json provides a version or description, those values take '
        'priority over the module-level constants. Use manifest.json as the single source of '
        'truth for packaged plugins.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 4: TOOL REGISTRATION
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('4. Tool Registration', level=1)

    doc.add_paragraph(
        'Plugins expose functionality to the AI Agent by defining a TOOLS list. '
        'Each entry follows the OpenAI Function Calling schema format. The Agent '
        'uses these schema definitions to decide when and how to invoke your tool.'
    )

    doc.add_heading('4.1 OpenAI Function Schema Format', level=2)

    doc.add_paragraph('Each tool entry has this structure:')
    add_code_block(doc, '{\n    "type": "function",\n    "function": {\n        "name": "my_tool_name",           # unique identifier\n        "description": "What this tool does",  # helps the AI decide when to call\n        "parameters": {\n            "type": "object",\n            "properties": {\n                "param_name": {\n                    "type": "string",     # string / integer / boolean / ...\n                    "description": "What this parameter is for"\n                }\n            },\n            "required": ["param_name"],    # list of required parameters\n            "additionalProperties": False\n        }\n    }\n}')

    doc.add_heading('Tool Naming Rules', level=2)
    naming_rules = [
        'Names must be unique across all plugins and built-in tools. Duplicate names cause a RuntimeError.',
        'Use snake_case naming: generate_uuid, code_review, save_note.',
        'Avoid names that conflict with the 14 built-in tools (read_file, write_file, exec_cmd, etc.).',
    ]
    for rule in naming_rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_heading('4.2 Complete TOOLS Example', level=2)
    doc.add_paragraph('This example registers a simple echo tool:')
    add_code_block(doc, 'TOOLS = [\n    {\n        "type": "function",\n        "function": {\n            "name": "echo",\n            "description": "Echoes back the input message. Useful for testing.",\n            "parameters": {\n                "type": "object",\n                "properties": {\n                    "message": {\n                        "type": "string",\n                        "description": "The message to echo back"\n                    },\n                    "repeat": {\n                        "type": "integer",\n                        "description": "Number of times to repeat (default 1)",\n                        "default": 1\n                    }\n                },\n                "required": ["message"],\n                "additionalProperties": False\n            }\n        }\n    }\n]')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 5: THE EXECUTE() FUNCTION
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('5. The execute() Function', level=1)

    doc.add_paragraph(
        'When the AI Agent decides to call one of your tools, the PluginManager dispatches '
        'the call to your plugin\'s execute() function. This is the core interface between '
        'the Agent and your plugin.'
    )

    doc.add_heading('5.1 Signature', level=2)
    add_code_block(doc, 'def execute(tool_name: str, args: dict, context: PluginContext) -> str:')

    doc.add_paragraph('Parameters:')
    params = [
        ('tool_name (str)', 'The name of the tool being invoked. Use this to dispatch calls when your plugin registers multiple tools.'),
        ('args (dict)', 'A dictionary of parameter values. Keys match the "properties" defined in your TOOLS schema. Required parameters are guaranteed to be present.'),
        ('context (PluginContext)', 'A read-only context object providing access to project root, config snapshot, storage, and a logger. See Chapter 7 for details.'),
    ]
    for name, desc in params:
        p = doc.add_paragraph()
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('5.2 Dispatching Pattern', level=2)
    doc.add_paragraph('Best practice for multi-tool plugins:')
    add_code_block(doc, 'def execute(tool_name: str, args: dict, context) -> str:\n    if tool_name == "my_tool_a":\n        return _handle_tool_a(args, context)\n    if tool_name == "my_tool_b":\n        return _handle_tool_b(args)\n    return f"Unknown tool: {tool_name}"')

    doc.add_heading('5.3 Return Value', level=2)
    doc.add_paragraph(
        'The execute() function must return a string. This string is passed back to the AI '
        'model as the tool call result and becomes part of the conversation context.'
    )
    return_tips = [
        'Use Markdown formatting for readability (the frontend renders Markdown).',
        'Keep responses concise. Very long outputs consume tokens.',
        'On error, return a descriptive message prefixed with \u274c (e.g., "\u274c File not found: path/to/file").',
        'Never raise unhandled exceptions. They are caught but produce ugly tracebacks.',
    ]
    for tip in return_tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 6: LIFECYCLE HOOKS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('6. Lifecycle Hooks', level=1)

    doc.add_paragraph(
        'Plugins can register callback functions for 15 lifecycle hooks spanning 4 layers. '
        'The PluginManager automatically discovers hooks by scanning the module for functions '
        'whose names match the hook list. Simply define a function with the correct name and '
        'signature, and it will be called automatically.'
    )

    doc.add_heading('6.1 Hook Categories', level=2)

    hook_table = doc.add_table(rows=1, cols=3)
    hook_table.style = 'Light Grid Accent 1'
    hook_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = hook_table.rows[0]
    for i, text in enumerate(['Layer', 'Hooks', 'Purpose']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    layer_data = [
        ('L1 \u2014 Lifecycle', 'on_agent_init, on_agent_shutdown', 'Plugin initialization and cleanup'),
        ('L2 \u2014 Task', 'on_task_start, on_task_done, on_task_error, on_task_stopped, on_task_timeout', 'Task-level events'),
        ('L3 \u2014 Step', 'before_step, after_step, before_tool_call, after_tool_call, on_user_input_required', 'Per-step and tool-call events'),
        ('L4 \u2014 Streaming', 'on_reasoning, on_content, on_event, on_usage_update', 'Real-time streaming events'),
    ]
    for row_data in layer_data:
        add_table_row(hook_table, row_data)

    doc.add_paragraph()

    doc.add_heading('6.2 L1 \u2014 Lifecycle Hooks', level=2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_agent_init(context)')
    run.bold = True
    doc.add_paragraph(
        'Called once when the Agent starts. Use this to initialize storage, load data files, '
        'or set up counters. This is the ideal place for one-time setup.'
    )
    add_code_block(doc, 'def on_agent_init(context):\n    context.storage["counter"] = 0\n    context.storage["session_start"] = time.time()\n    context.logger.info("My plugin loaded!")')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_agent_shutdown(context)')
    run.bold = True
    doc.add_paragraph(
        'Called once when the Agent shuts down. Use this to persist state, write summary '
        'statistics, or clean up resources.'
    )
    add_code_block(doc, 'def on_agent_shutdown(context):\n    total = context.storage.get("counter", 0)\n    context.logger.info(f"Shutting down. Total operations: {total}")')

    doc.add_heading('6.3 L2 \u2014 Task Hooks', level=2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_task_start(task_text: str, context)')
    run.bold = True
    doc.add_paragraph(
        'Called when the user submits a new task. Receives the full task text. '
        'Use this to detect task keywords and prepare context.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_task_done(summary: str, final_reply: str, context)')
    run.bold = True
    doc.add_paragraph(
        'Called when a task completes successfully. Receives the task summary and the '
        'final AI reply. Use this to auto-save notes or record completion statistics.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_task_error(error_msg: str, context)')
    run.bold = True
    doc.add_paragraph(
        'Called when a task fails. Receives the error message. Use this to log errors '
        'or trigger recovery actions.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_task_stopped(context)')
    run.bold = True
    doc.add_paragraph('Called when the user manually stops a running task.')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_task_timeout(elapsed: float, context)')
    run.bold = True
    doc.add_paragraph(
        'Called when a task exceeds the configured timeout. Receives the elapsed seconds. '
        'Take care: this hook itself has a 5-second hard timeout (HOOK_TIMEOUT).'
    )

    doc.add_heading('6.4 L3 \u2014 Step Hooks (Mutating)', level=2)

    doc.add_paragraph(
        'These hooks can modify the data flow. Their return values are used by the Agent.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('before_step(step: int, messages: list, context) \u2192 list | None')
    run.bold = True
    doc.add_paragraph(
        'Called before each ReAct reasoning step. Receives the current step number and '
        'the messages list being sent to the model. Return a modified list to alter '
        'the conversation context, or None to make no changes.'
    )
    add_code_block(doc, 'def before_step(step: int, messages: list, context):\n    """Inject a system reminder every 10 steps."""\n    if step % 10 == 0:\n        messages.append({"role": "system", "content": "Remember to save progress!"})\n    return messages  # return modified list')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('after_step(step: int, reasoning: str, content: str, tool_calls: list, context)')
    run.bold = True
    doc.add_paragraph(
        'Called after each ReAct step completes. Non-mutating \u2014 return value is ignored. '
        'Use this for analytics or logging.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('before_tool_call(tool_name: str, args: dict, context) \u2192 dict | None')
    run.bold = True
    doc.add_paragraph(
        'Called right before any tool (built-in or plugin) executes. '
        'Return a modified args dict to alter the parameters, or None to allow the call to proceed unchanged. '
        'Important: If no plugin registers this hook, the tool call is never blocked \u2014 '
        'the short-circuit optimizes for the common case.'
    )
    add_code_block(doc, 'def before_tool_call(tool_name: str, args: dict, context):\n    """Log and time every tool call."""\n    context.storage["_tool_start"] = time.time()\n    context.logger.info(f"Tool call: {tool_name}")\n    return args  # pass through unchanged')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('after_tool_call(tool_name: str, args: dict, result: str, context) \u2192 str | None')
    run.bold = True
    doc.add_paragraph(
        'Called after a tool completes. Receives the original args and the result string. '
        'Return a modified result string to alter what the AI sees, or None to leave it unchanged.'
    )
    add_code_block(doc, 'def after_tool_call(tool_name: str, args: dict, result: str, context):\n    """Anonymize sensitive data in tool results."""\n    result = result.replace("API_KEY_12345", "***REDACTED***")\n    return result')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('on_user_input_required(question: str, context)')
    run.bold = True
    doc.add_paragraph(
        'Called when the Agent pauses and waits for user input (via ask_user). '
        'Use this to notify external systems or log pending questions.'
    )

    doc.add_heading('6.5 L4 \u2014 Streaming Event Hooks', level=2)

    stream_hooks = [
        ('on_reasoning(token: str, context)', 'Called for each reasoning token (the AI\'s "inner monologue"). Fires frequently.'),
        ('on_content(token: str, context)', 'Called for each content token in the AI\'s final response. Fires very frequently \u2014 keep handlers lightweight.'),
        ('on_event(event_type: str, data: str, context)', 'Called for structured events. event_type is a single character prefix (T, R, C, Q, D, E, U, F). See the Event Stream Protocol.'),
        ('on_usage_update(usage: dict, context)', 'Called when token usage statistics are updated. usage contains input_tokens, output_tokens, tool_call_tokens.'),
    ]
    for name, desc in stream_hooks:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        doc.add_paragraph(desc)

    doc.add_heading('6.6 Mutating vs Non-Mutating Hooks', level=2)
    doc.add_paragraph(
        'Three hooks can mutate the data flow: before_step, before_tool_call, and after_tool_call. '
        'For these hooks, the first non-None return value from any plugin wins. '
        'All other hooks are fire-and-forget \u2014 their return values are ignored.'
    )
    doc.add_paragraph(
        'Hook timeout: Every hook callback has a hard timeout of 5 seconds (HOOK_TIMEOUT). '
        'If a hook takes longer, the calling thread is abandoned (not joined) to prevent '
        'blocking the Agent. Keep hook implementations fast and non-blocking.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 7: PLUGINCONTEXT API
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('7. PluginContext API', level=1)

    doc.add_paragraph(
        'Every hook and the execute() function receive a PluginContext object. '
        'This is your window into the Agent\'s environment. The context is read-only '
        'except for the storage dict and the logger.'
    )

    doc.add_heading('7.1 Attributes', level=2)

    ctx_table = doc.add_table(rows=1, cols=3)
    ctx_table.style = 'Light Grid Accent 1'
    ctx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ctx_table.rows[0]
    for i, text in enumerate(['Attribute', 'Type', 'Description']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    ctx_data = [
        ('project_root', 'str', 'Absolute path of the current workspace / project root. All file operations should be scoped under this path.'),
        ('app_dir', 'str', 'Application data directory (config, memories, logs). Use for plugin-private data files.'),
        ('config', 'dict', 'Read-only snapshot of the current config.json. Updated before each hook call.'),
        ('storage', 'dict', 'Per-plugin key-value store. Survives across hooks within one Agent session. Cleared on restart.'),
        ('logger', 'SimpleLogger', 'A pre-configured logger that writes to both stdout and a plugin.log file in app_dir.'),
        ('current_step', 'int', 'The current ReAct step number. Updated before each step hook.'),
        ('total_usage', 'dict', 'Cumulative token usage: {"input_tokens": N, "output_tokens": N, "tool_call_tokens": N}.'),
    ]
    for row_data in ctx_data:
        add_table_row(ctx_table, row_data)

    doc.add_paragraph()
    doc.add_paragraph('Using storage:')
    add_code_block(doc, '# Store arbitrary data (survives across hooks)\ncontext.storage["my_counter"] = context.storage.get("my_counter", 0) + 1\n\n# Read config\nmax_steps = context.config.get("max_steps", 128)\n\n# Build a safe file path\nfull_path = os.path.join(context.project_root, "relative/path.txt")\nfull_path = os.path.normpath(full_path)')

    doc.add_heading('7.2 SimpleLogger', level=2)
    doc.add_paragraph(
        'Each plugin gets a SimpleLogger instance pre-configured with the plugin name. '
        'Messages are printed to stdout and appended to plugin.log in the app directory.'
    )

    add_code_block(doc, '# Available methods:\ncontext.logger.info("Informational message")\ncontext.logger.warn("Warning message")\ncontext.logger.error("Error message")\ncontext.logger.debug("Debug message")\n\n# Output format:\n# [2026-08-01 17:32:33] [INFO] [My Plugin] Informational message')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 8: SECURITY SYSTEM
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('8. Security System', level=1)

    doc.add_paragraph(
        'Vibe Coding Agent employs a multi-layer security system to protect the host '
        'environment from potentially malicious or buggy plugins. All security features '
        'are individually togglable via config.json keys.'
    )

    doc.add_heading('8.1 AST Source Audit', level=2)
    doc.add_paragraph(
        'Before a plugin module is loaded, its source code is parsed into an AST '
        '(Abstract Syntax Tree) and walked to detect dangerous patterns. The audit '
        'checks for:'
    )

    audit_checks = [
        'Shell/process execution: os.system(), subprocess.run(), etc. (CRITICAL)',
        'Code execution: eval(), exec(), compile(), __import__() (CRITICAL)',
        'Native code loading: ctypes, cffi (CRITICAL)',
        'File deletion: os.remove(), shutil.rmtree() (WARNING)',
        'Network access: socket, requests, urllib (WARNING)',
        'Deserialization: pickle, marshal, yaml.unsafe_load() (WARNING)',
        'System manipulation: sys.modules, sys.setprofile(), builtins override (WARNING)',
        'Process termination: os._exit(), os.kill() (CRITICAL)',
    ]
    for check in audit_checks:
        doc.add_paragraph(check, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'The audit_level configuration controls behavior:\n'
        '\u2022 "off" \u2014 Skip source audit entirely.\n'
        '\u2022 "warn" (default) \u2014 Log warnings but allow loading.\n'
        '\u2022 "block" \u2014 Reject plugins with any CRITICAL findings.'
    )

    doc.add_heading('8.2 Import Restriction', level=2)
    doc.add_paragraph(
        'After the source audit, import blockers are installed in sys.meta_path to '
        'prevent plugin modules from importing dangerous modules at runtime. The blockers '
        'use inspect.stack() to only restrict imports originating from plugin code '
        '(modules with names starting with vibe_plugin_).'
    )

    doc.add_paragraph('Three restriction levels:')
    doc.add_paragraph('\u2022 "off" (default) \u2014 No import restrictions.')
    doc.add_paragraph(
        '\u2022 "safe" \u2014 Block all known dangerous modules (subprocess, ctypes, socket, pickle, '
        'marshal, telnetlib, ftplib, smtplib).'
    )
    doc.add_paragraph(
        '\u2022 "strict" \u2014 Only allow a whitelist of safe modules (json, re, datetime, math, '
        'random, collections, itertools, pathlib, etc.). All other imports are blocked for plugin code.'
    )

    doc.add_heading('8.3 Permission Declaration', level=2)
    doc.add_paragraph(
        'When require_permissions is enabled in config.json, plugins using manifest.json '
        'must declare required permissions. The security system cross-references declared '
        'permissions with audit findings and blocks loading if there is a mismatch.'
    )

    doc.add_paragraph('Supported permission values:')
    perms = [
        '"process" \u2014 Required for plugins that spawn processes or execute code.',
        '"network" \u2014 Required for plugins that access the network.',
        '"file_write" \u2014 Required for plugins that delete or move files.',
        '"file_read" \u2014 Required for plugins that read files.',
    ]
    for perm in perms:
        doc.add_paragraph(perm, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Example manifest.json permission declaration:')
    add_code_block(doc, '{\n    "name": "Network Monitor",\n    "permissions": ["network", "file_read"]\n}')

    doc.add_heading('8.4 Resource Limits', level=2)
    doc.add_paragraph(
        'When resource_limit is enabled, each plugin module load is wrapped with '
        'ResourceLimiter which enforces:'
    )
    doc.add_paragraph('\u2022 CPU time limit: 30 seconds (Unix: signal.SIGALRM, Windows: Timer thread)')
    doc.add_paragraph('\u2022 Memory limit: 512 MB (Unix: resource.setrlimit RLIMIT_AS)')
    doc.add_paragraph(
        'These limits only apply during module loading. Once loaded, hook execution '
        'already has the 5-second HOOK_TIMEOUT protection.'
    )

    doc.add_heading('8.5 Security Configuration Summary', level=2)

    sec_table = doc.add_table(rows=1, cols=3)
    sec_table.style = 'Light Grid Accent 1'
    sec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = sec_table.rows[0]
    for i, text in enumerate(['Config Key', 'Default', 'Values']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    sec_data = [
        ('plugin_security_audit', '"warn"', '"off" | "warn" | "block"'),
        ('plugin_security_import_restrict', '"off"', '"off" | "safe" | "strict"'),
        ('plugin_security_require_permissions', 'false', 'true | false'),
        ('plugin_security_resource_limit', 'false', 'true | false'),
    ]
    for row_data in sec_data:
        add_table_row(sec_table, row_data)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 9: MANIFEST.JSON REFERENCE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('9. manifest.json Reference', level=1)

    doc.add_paragraph(
        'The manifest.json file is required for package-style plugins. It provides '
        'metadata and configuration that is read before the Python module is loaded.'
    )

    doc.add_heading('Complete Schema', level=2)
    add_code_block(doc, '{\n    "name": "My Plugin",\n    "version": "1.0.0",\n    "publisher": "Your Name",\n    "author": "Your Name",          // alias for "publisher"\n    "description": "A concise description of what the plugin does.",\n    "entry": "plugin.py",           // entry-point file (default: "plugin.py")\n    "enabled": true,                // can be false to disable without deleting\n    "permissions": [                // required permissions (see \u00a78.3)\n        "file_read",\n        "file_write"\n    ]\n}')

    doc.add_heading('Field Reference', level=2)

    mf_table = doc.add_table(rows=1, cols=4)
    mf_table.style = 'Light Grid Accent 1'
    mf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = mf_table.rows[0]
    for i, text in enumerate(['Field', 'Type', 'Required', 'Description']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    mf_data = [
        ('name', 'string', 'Yes', 'Plugin display name. Takes priority over PLUGIN_NAME constant.'),
        ('version', 'string', 'No', 'Semantic version. Takes priority over PLUGIN_VERSION.'),
        ('publisher', 'string', 'No', 'Publisher name. "author" is an accepted alias.'),
        ('description', 'string', 'No', 'Plugin description. Takes priority over PLUGIN_DESCRIPTION.'),
        ('entry', 'string', 'No', 'Entry Python file relative to the plugin directory. Default: "plugin.py".'),
        ('enabled', 'boolean', 'No', 'Whether the plugin is enabled. Default: true. Set false to disable.'),
        ('permissions', 'string[]', 'No', 'List of permission strings. Only checked when require_permissions is enabled.'),
    ]
    for row_data in mf_data:
        add_table_row(mf_table, row_data)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 10: STEP-BY-STEP TUTORIAL
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('10. Step-by-Step Tutorial', level=1)

    doc.add_paragraph(
        'This tutorial walks you through creating a complete plugin from scratch. '
        'We\'ll build a "Weather Logger" plugin that records weather observations '
        'and generates simple summaries.'
    )

    doc.add_heading('10.1 Create the File', level=2)
    doc.add_paragraph(
        'Create a new file named weather_logger.py in one of your plugin directories. '
        'You can find your plugin directories in config.json under plugin_dirs.'
    )
    add_code_block(doc, '# weather_logger.py\n# A simple plugin for logging weather observations')

    doc.add_heading('10.2 Declare Metadata', level=2)
    doc.add_paragraph('Add the required module-level constants:')
    add_code_block(doc, 'PLUGIN_NAME = "Weather Logger"\nPLUGIN_PUBLISHER = "Your Name"\nPLUGIN_VERSION = "1.0.0"\nPLUGIN_DESCRIPTION = "Logs weather observations and generates summaries."')

    doc.add_heading('10.3 Register a Tool', level=2)
    doc.add_paragraph('Define the TOOLS list with an OpenAI function schema:')
    add_code_block(doc, 'TOOLS = [\n    {\n        "type": "function",\n        "function": {\n            "name": "log_weather",\n            "description": "Record a weather observation with temperature and conditions.",\n            "parameters": {\n                "type": "object",\n                "properties": {\n                    "location": {\n                        "type": "string",\n                        "description": "City or location name"\n                    },\n                    "temperature_c": {\n                        "type": "number",\n                        "description": "Temperature in Celsius"\n                    },\n                    "conditions": {\n                        "type": "string",\n                        "description": "Weather conditions, e.g. sunny, rainy, cloudy"\n                    }\n                },\n                "required": ["location", "temperature_c"],\n                "additionalProperties": False\n            }\n        }\n    }\n]')

    doc.add_heading('10.4 Implement execute()', level=2)
    doc.add_paragraph('Write the execute() function to handle the tool call:')
    add_code_block(doc, 'def execute(tool_name: str, args: dict, context) -> str:\n    if tool_name == "log_weather":\n        location = args["location"]\n        temp = args["temperature_c"]\n        conditions = args.get("conditions", "unknown")\n\n        # Store the observation\n        obs = context.storage.get("observations", [])\n        obs.append({\n            "location": location,\n            "temp": temp,\n            "conditions": conditions,\n            "time": __import__("datetime").datetime.now().isoformat()\n        })\n        context.storage["observations"] = obs\n\n        context.logger.info(f"Weather logged: {location} {temp}C {conditions}")\n\n        return (\n            f"[OK] Weather recorded for **{location}**:\\n"\n            f"  Temperature: {temp}C\\n"\n            f"  Conditions: {conditions}\\n"\n            f"  Total observations: {len(obs)}"\n        )\n\n    return f"Unknown tool: {tool_name}"')

    doc.add_heading('10.5 Add Lifecycle Hooks', level=2)
    doc.add_paragraph('Add hooks for initialization and task completion:')
    add_code_block(doc, 'def on_agent_init(context):\n    """Initialize the observation log."""\n    context.storage["observations"] = []\n    context.logger.info("Weather Logger ready! \u2600\ufe0f")\n\n\ndef on_task_done(summary: str, final_reply: str, context):\n    """Generate a summary when tasks complete."""\n    obs = context.storage.get("observations", [])\n    if not obs:\n        return\n\n    latest = obs[-1]\n    context.logger.info(\n        f"Task done. Last weather: {latest[\'location\']} "\n        f"{latest[\'temp\']}\u00b0C ({latest[\'conditions\']})"\n    )')

    doc.add_heading('10.6 Test Your Plugin', level=2)
    doc.add_paragraph('After saving the file:')
    test_steps = [
        'Restart the Vibe Coding Agent application.',
        'Open Settings and verify your plugin appears in the plugin list.',
        'Check that no security audit errors are shown.',
        'Try a task like: "Log the weather in Tokyo: 22\u00b0C, partly cloudy".',
        'The Agent should invoke your log_weather tool and display the result.',
    ]
    for i, step in enumerate(test_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 11: BEST PRACTICES
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('11. Best Practices', level=1)

    practices = [
        ('Keep Hooks Lightweight',
         'Hook callbacks have a 5-second timeout. Avoid blocking I/O, long computations, '
         'or infinite loops. If you need to do heavy work, consider spawning a background '
         'thread (but be aware of the security implications).'),
        ('Use context.storage for State',
         'The storage dict is your per-plugin state container. It survives across all hooks '
         'within a session. Initialize default values in on_agent_init.'),
        ('Validate Tool Arguments',
         'Even though the schema defines types, always validate critical arguments '
         '(file paths exist, numbers are in range, strings are not empty). '
         'The AI can occasionally pass unexpected values.'),
        ('Scope File Operations',
         'Always build file paths using os.path.join(context.project_root, relative_path) '
         'followed by os.path.normpath(). Never access files outside the project root.'),
        ('Use Descriptive Tool Descriptions',
         'The AI model uses the tool description to decide which tool to call. '
         'Write clear, specific descriptions that include when the tool should and '
         'should not be used.'),
        ('Handle Errors Gracefully',
         'Catch exceptions in execute() and return user-friendly error messages. '
         'Never let exceptions propagate \u2014 they result in ugly tracebacks in the conversation.'),
        ('Log Strategically',
         'Use context.logger.info() for significant events and context.logger.debug() '
         'for detailed tracing. This makes debugging easier without flooding the logs.'),
        ('Test with Security Enabled',
         'Test your plugin with plugin_security_audit set to "block" to ensure it passes '
         'all security checks. This prevents surprises when users enable strict security.'),
        ('Document Your Plugin',
         'Include a clear PLUGIN_DESCRIPTION. For package plugins, add a README.md '
         'alongside manifest.json. Users will appreciate good documentation.'),
        ('Watch for Tool Name Conflicts',
         'Choose unique tool names. Check the 14 built-in tools and other installed plugins '
         'to avoid name collisions, which cause RuntimeError during loading.'),
    ]

    for title, desc in practices:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 12: TROUBLESHOOTING
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('12. Troubleshooting', level=1)

    issues = [
        ('Plugin not appearing in the UI',
         '\u2022 Verify the plugin directory is listed in config.json \u2192 plugin_dirs.\n'
         '\u2022 Check that the file is not named __init__.py.\n'
         '\u2022 For package plugins, ensure manifest.json exists and is valid JSON.\n'
         '\u2022 Check that PLUGIN_NAME is defined in the module.'),
        ('"Security audit blocked" error',
         '\u2022 Your plugin uses dangerous patterns (e.g., os.system, subprocess).\n'
         '\u2022 Temporarily set plugin_security_audit to "warn" to see the full audit report.\n'
         '\u2022 Refactor to use safe alternatives (e.g., use context.logger instead of subprocess).\n'
         '\u2022 If the pattern is necessary, add appropriate permissions to manifest.json.'),
        ('"Import blocked" error',
         '\u2022 Your plugin imports a restricted module (subprocess, ctypes, socket, etc.).\n'
         '\u2022 Set plugin_security_import_restrict to "off" if you trust the plugin source.\n'
         '\u2022 For strict mode, add your module to STRICT_SAFE_MODULES in security.py.'),
        ('Tool call returns "Unknown tool"',
         '\u2022 The execute() function does not handle the tool_name being dispatched.\n'
         '\u2022 Check for typos in the tool name string comparison.\n'
         '\u2022 Verify the tool is registered in the TOOLS list with the correct name.'),
        ('Hook not being called',
         '\u2022 Ensure the function name exactly matches one of the 15 hook names.\n'
         '\u2022 Check the function signature \u2014 wrong parameter count causes silent failure.\n'
         '\u2022 Verify the plugin loaded successfully (check plugin list in Settings).'),
        ('Plugin crashes the Agent',
         '\u2022 Unhandled exceptions in hooks are silently caught.\n'
         '\u2022 Exceptions in execute() are caught and returned as error strings.\n'
         '\u2022 Check plugin.log in the app directory for error traces.\n'
         '\u2022 Use try/except blocks around risky operations.'),
    ]

    for title, solution in issues:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        doc.add_paragraph(solution)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CHAPTER 13: APPENDIX
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('13. Appendix', level=1)

    doc.add_heading('13.1 Complete Hook Reference', level=2)

    hook_ref = doc.add_table(rows=1, cols=4)
    hook_ref.style = 'Light Grid Accent 1'
    hook_ref.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = hook_ref.rows[0]
    for i, text in enumerate(['Hook', 'Signature', 'Mutating', 'Layer']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    hooks_data = [
        ('on_agent_init', '(context)', 'No', 'L1 Lifecycle'),
        ('on_agent_shutdown', '(context)', 'No', 'L1 Lifecycle'),
        ('on_task_start', '(task_text, context)', 'No', 'L2 Task'),
        ('on_task_done', '(summary, final_reply, context)', 'No', 'L2 Task'),
        ('on_task_error', '(error_msg, context)', 'No', 'L2 Task'),
        ('on_task_stopped', '(context)', 'No', 'L2 Task'),
        ('on_task_timeout', '(elapsed, context)', 'No', 'L2 Task'),
        ('before_step', '(step, messages, context) \u2192 list', 'Yes', 'L3 Step'),
        ('after_step', '(step, reasoning, content, tool_calls, context)', 'No', 'L3 Step'),
        ('before_tool_call', '(tool_name, args, context) \u2192 dict', 'Yes', 'L3 Step'),
        ('after_tool_call', '(tool_name, args, result, context) \u2192 str', 'Yes', 'L3 Step'),
        ('on_user_input_required', '(question, context)', 'No', 'L3 Step'),
        ('on_reasoning', '(token, context)', 'No', 'L4 Streaming'),
        ('on_content', '(token, context)', 'No', 'L4 Streaming'),
        ('on_event', '(event_type, data, context)', 'No', 'L4 Streaming'),
        ('on_usage_update', '(usage, context)', 'No', 'L4 Streaming'),
    ]
    for row_data in hooks_data:
        add_table_row(hook_ref, row_data)

    doc.add_paragraph()

    doc.add_heading('13.2 Dangerous Pattern Registry', level=2)

    doc.add_paragraph(
        'The following patterns are flagged by the AST source audit. Plugins using '
        'these will trigger warnings or be blocked depending on the audit_level.'
    )

    danger_table = doc.add_table(rows=1, cols=3)
    danger_table.style = 'Light Grid Accent 1'
    danger_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = danger_table.rows[0]
    for i, text in enumerate(['Pattern', 'Severity', 'Category']):
        hdr.cells[i].text = text
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    danger_data = [
        ('os.system(), os.popen()', 'CRITICAL', 'shell_exec'),
        ('subprocess.call(), run(), Popen(), check_output()', 'CRITICAL', 'shell_exec'),
        ('eval(), exec(), compile()', 'CRITICAL', 'code_exec'),
        ('__import__(), importlib.import_module()', 'CRITICAL', 'import_bypass'),
        ('ctypes, cffi', 'CRITICAL', 'native_exec'),
        ('os._exit(), os.kill()', 'CRITICAL', 'process_terminate'),
        ('os.remove(), os.unlink(), os.rmdir()', 'WARNING', 'file_delete'),
        ('shutil.rmtree(), shutil.move()', 'WARNING', 'file_delete / file_move'),
        ('socket, http, urllib, requests', 'WARNING', 'network'),
        ('ftplib, smtplib, poplib, telnetlib', 'WARNING', 'network'),
        ('pickle, marshal, yaml.unsafe_load()', 'WARNING', 'deserialization'),
        ('sys.modules, sys.setprofile(), sys.settrace()', 'WARNING', 'sys_manipulation'),
        ('builtins override, sys.exit()', 'WARNING', 'builtin_override / sys_manipulation'),
    ]
    for row_data in danger_data:
        add_table_row(danger_table, row_data)

    doc.add_paragraph()

    doc.add_heading('13.3 PluginContext Quick Reference', level=2)

    add_code_block(doc, '# PluginContext attributes\ncontext.project_root     # str  \u2014 workspace root path\ncontext.app_dir          # str  \u2014 app data directory\ncontext.config           # dict \u2014 config.json snapshot (read-only)\ncontext.storage          # dict \u2014 per-plugin key-value store\ncontext.logger           # SimpleLogger \u2014 .info() .warn() .error() .debug()\ncontext.current_step     # int  \u2014 current ReAct step\ncontext.total_usage      # dict \u2014 {"input_tokens", "output_tokens", "tool_call_tokens"}\n\n# Safe file path construction\nimport os\nfull_path = os.path.normpath(os.path.join(context.project_root, rel_path))')

    doc.add_paragraph()
    doc.add_heading('13.4 Official Plugins Source Reference', level=2)
    doc.add_paragraph(
        'Study the official plugins for complete, working examples:\n\n'
        '\u2022 official_plugins/code_reviewer.py \u2014 AST-based code quality review\n'
        '\u2022 official_plugins/dev_utilities.py \u2014 UUID, password, hash, timestamp, line count\n'
        '\u2022 official_plugins/note_manager.py \u2014 Persistent notes with tags and full-text search\n'
        '\u2022 official_plugins/time_tracker.py \u2014 Task timing, tool call profiling, efficiency reports'
    )

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2500\u2500\u2500 End of Document \u2500\u2500\u2500')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.datetime.now().strftime("%B %d, %Y")}  |  NORP Vibe Coding Agent v1.0')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    return doc


if __name__ == '__main__':
    print('Creating plugin_develop_guide.docx ...')
    doc = create_document()
    doc.save('plugin_develop_guide.docx')
    print('Done! Output: plugin_develop_guide.docx')
