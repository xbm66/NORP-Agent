"""
翻译 whats_new.docx 为繁体中文、英语、俄语，生成三份新 docx。
保留原文档格式（段落样式、表格结构、run级别格式）。
使用 MyMemory 翻译 API 进行英语和俄语翻译。
"""
import time
import docx
from docx import Document
from docx.oxml.ns import qn
from zhconv import convert as zh_convert
from deep_translator import MyMemoryTranslator


def translate_zh_tw(text: str) -> str:
    """简体中文 → 繁体中文（使用 zhconv 字符转换）"""
    if not text.strip():
        return text
    return zh_convert(text, 'zh-tw')


def translate_en(text: str) -> str:
    """中文 → 英语（使用 MyMemory）"""
    if not text.strip():
        return text
    try:
        result = MyMemoryTranslator(source='zh-CN', target='en-GB').translate(text)
        return result if result else text
    except Exception as e:
        print(f"    [EN] Error: {e}, retrying with auto-detect...")
        try:
            result = MyMemoryTranslator(source='auto', target='en-GB').translate(text)
            return result if result else text
        except Exception as e2:
            print(f"    [EN] Retry failed: {e2}")
            return text


def translate_ru(text: str) -> str:
    """中文 → 俄语（使用 MyMemory，两步：中文→英语→俄语以获得更好质量）"""
    if not text.strip():
        return text
    try:
        # 直接中文→俄语
        result = MyMemoryTranslator(source='zh-CN', target='ru-RU').translate(text)
        # 如果结果看起来是英文（没翻译成功），则通过英语中转
        if result and all(ord(c) < 128 for c in result.replace(' ', '')):
            # 可能是英文，需要两步翻译
            en = MyMemoryTranslator(source='zh-CN', target='en-GB').translate(text)
            result = MyMemoryTranslator(source='en-GB', target='ru-RU').translate(en)
        return result if result else text
    except Exception as e:
        print(f"    [RU] Error: {e}")
        return text


def copy_paragraph_format(src_para, dst_para):
    """复制段落格式"""
    dst_para.style = src_para.style
    dst_para.alignment = src_para.alignment


def copy_run_format(src_run, dst_run):
    """复制 run 级别的格式"""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb


def translate_document(src_path: str, dst_path: str, translator_func, lang_label: str):
    """翻译整个文档并保存"""
    print(f"\n{'='*60}")
    print(f"Translating to {lang_label}...")
    print(f"Source: {src_path}")
    print(f"Output: {dst_path}")

    src_doc = Document(src_path)
    dst_doc = Document()

    # 复制页面设置
    for i, section in enumerate(src_doc.sections):
        if i >= len(dst_doc.sections):
            dst_doc.add_section()
        dst_section = dst_doc.sections[i]
        dst_section.page_width = section.page_width
        dst_section.page_height = section.page_height
        dst_section.left_margin = section.left_margin
        dst_section.right_margin = section.right_margin
        dst_section.top_margin = section.top_margin
        dst_section.bottom_margin = section.bottom_margin

    # 处理段落
    total_paras = len(src_doc.paragraphs)
    for idx, para in enumerate(src_doc.paragraphs):
        if idx % 10 == 0:
            print(f"  Paragraphs: {idx}/{total_paras}")

        text = para.text.strip()
        if not text and not para.runs:
            dst_doc.add_paragraph()
            continue

        # 翻译
        translated = translator_func(text) if text else ""

        # 创建新段落
        new_para = dst_doc.add_paragraph()
        copy_paragraph_format(para, new_para)

        # 添加翻译后的文本，保留第一个 run 的格式
        if para.runs:
            first_run = para.runs[0]
            new_run = new_para.add_run(translated)
            copy_run_format(first_run, new_run)
        else:
            new_para.add_run(translated)

        # 延迟避免 API 限流
        if idx % 5 == 0 and idx > 0:
            time.sleep(0.3)

    # 处理表格
    print(f"  Processing {len(src_doc.tables)} tables...")
    for ti, table in enumerate(src_doc.tables):
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        new_table = dst_doc.add_table(rows=num_rows, cols=num_cols)

        if table.style:
            new_table.style = table.style

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                src_cell_text = cell.text
                translated = translator_func(src_cell_text) if src_cell_text.strip() else src_cell_text
                new_table.rows[ri].cells[ci].text = translated
                if cell.paragraphs:
                    src_cell_para = cell.paragraphs[0]
                    dst_cell_para = new_table.rows[ri].cells[ci].paragraphs[0]
                    copy_paragraph_format(src_cell_para, dst_cell_para)
                time.sleep(0.1)

    # 保存
    dst_doc.save(dst_path)
    print(f"  Saved: {dst_path}")


def main():
    src = "whats_new.docx"

    # 1. 繁体中文（本地转换，速度很快）
    translate_document(src, "whats_new_zh-tw.docx", translate_zh_tw, "繁體中文")

    # 2. 英语
    translate_document(src, "whats_new_en.docx", translate_en, "English")

    # 3. 俄语
    translate_document(src, "whats_new_ru.docx", translate_ru, "Русский")

    print("\n" + "="*60)
    print("All translations completed!")
    print("  - whats_new_zh-tw.docx (繁體中文)")
    print("  - whats_new_en.docx (English)")
    print("  - whats_new_ru.docx (Русский)")


if __name__ == "__main__":
    main()
